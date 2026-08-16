"""Gera um relatorio Word a partir de um ficheiro Excel e de um plano em JSON.

Fluxo: le o .xlsx -> valida o plano -> verifica os dados -> desenha os graficos
em PNG -> monta o .docx.

A verificacao corre sempre. Se encontrar avisos, a geracao e bloqueada ate o
utilizador passar --avisos-aceites. Isto e mecanico de proposito: uma promessa
em prosa nao trava nada.

O ficheiro Excel de origem nunca e escrito nem alterado.
"""

from __future__ import annotations

import argparse
import csv
import datetime
import io
import json
import math
import re
import shutil
import sys
import tempfile
import unicodedata
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # sem janela; obrigatorio para correr sem ecra

import matplotlib.ticker
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TIPOS_GRAFICO = ("barras", "barras_horizontais", "linhas", "area", "circular",
                 "dispersao")
# a dispersao nao agrupa nada: cruza duas colunas de numeros, um ponto por linha
TIPOS_SEM_AGREGACAO = ("dispersao",)

# forca da correlacao, com o criterio escrito ao lado como todos os outros
R_CORRELACAO_FRACA = 0.3
R_CORRELACAO_FORTE = 0.7
AGREGACOES = ("soma", "media", "contagem")

CHAVES_OBRIGATORIAS = ("tipo", "eixo_x", "titulo")
CHAVES_OPCIONAIS = ("folha", "eixo_y", "nota", "tabela_dados", "linha_cabecalho",
                    "coluna_periodo", "eixo_temporal", "previsao", "meta",
                    "serie", "agregacao", "filtro")

AMBITOS_META = ("total", "categoria", "serie")
# acima disto as linhas confundem-se umas com as outras
MAX_SERIES = 6
# 24 periodos sem UMA unica descida e raro em dados verdadeiros; 12 nao e.
# Com o limite em 12, uma serie legitima a crescer todos os meses era acusada
# de ser um acumulado -- falso positivo apanhado pelos proprios testes.
MIN_PERIODOS_ACUMULADO = 24

# ate onde se procura o cabecalho quando a tabela nao comeca na primeira linha
MAX_LINHAS_PROCURA_CABECALHO = 25
# nomes de coluna listados numa mensagem de erro; 93 numa linha nao se leem
MAX_COLUNAS_LISTADAS = 12

MAX_LINHAS_TABELA = 30
MAX_FATIAS_CIRCULAR = 12
# acima disto as barras ficam finas como cabelos e os rotulos ilegiveis
MAX_CATEGORIAS_LEGIVEIS = 40
# anos aceites: dados reais recuam a 1850 e mais atras
ANO_MINIMO = 1000
ANO_MAXIMO = 2999

# rotulos que denunciam uma linha de totais deixada por um sistema de exportacao
ROTULOS_DE_TOTAL = frozenset({"total", "totais", "total geral", "subtotal", "soma", "sum"})

# simbolos a ignorar quando os numeros vem gravados como texto
SIMBOLOS_MOEDA = ("€", "$", "£", "¥", "R$", "EUR", "USD")

# Minimos para cada metodo. Aplicar uma regressao a tres pontos faz um relatorio
# PARECER serio sendo falso; por isso, abaixo destes numeros o metodo nao corre e
# o relatorio diz que nao correu e porque.
MIN_PERIODOS_REGRESSAO = 4
MIN_PERIODOS_CRESCIMENTO = 3
MIN_CATEGORIAS_ATIPICOS = 5
MIN_CICLOS_SAZONALIDADE = 2

# Guardas da previsao. Uma previsao e a afirmacao mais facil de fazer e a mais
# dificil de defender: sem estes minimos, uma reta forcada sobre pontos
# dispersos daria numeros com ar de rigor e sem rigor nenhum.
MIN_PERIODOS_PREVISAO = 8
R2_MINIMO_PREVISAO = 0.3
MAX_HORIZONTE_PREVISAO = 12
FRACAO_MAXIMA_HORIZONTE = 1 / 3
CONFIANCA_PREVISAO = 95

# t de Student a 95% (bilateral), por graus de liberdade. Acima de 30 usa-se
# 1,96, que e o limite. Embutida para nao trazer o scipy so por causa disto.
T_95 = {
    1: 12.706, 2: 4.303, 3: 3.182, 4: 2.776, 5: 2.571, 6: 2.447, 7: 2.365,
    8: 2.306, 9: 2.262, 10: 2.228, 11: 2.201, 12: 2.179, 13: 2.160, 14: 2.145,
    15: 2.131, 16: 2.120, 17: 2.110, 18: 2.101, 19: 2.093, 20: 2.086,
    21: 2.080, 22: 2.074, 23: 2.069, 24: 2.064, 25: 2.060, 26: 2.056,
    27: 2.052, 28: 2.048, 29: 2.045, 30: 2.042,
}

# Limiares dos termos qualitativos. Cada um aparece no relatorio com o numero E o
# criterio ao lado: nenhuma palavra qualitativa entra sem regra a sustenta-la.
CV_DISPERSAO_BAIXA = 15.0
CV_DISPERSAO_ELEVADA = 35.0
R2_AJUSTE_FRACO = 0.3
R2_AJUSTE_FORTE = 0.7
PESO_CONCENTRACAO_BAIXA = 25.0
PESO_CONCENTRACAO_ELEVADA = 50.0

MESES_PT = (
    "janeiro", "fevereiro", "marco", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
)
LARGURA_IMAGEM = Inches(6.0)

# tamanhos pensados para a figura ja encolhida dentro da pagina do Word
plt.rcParams.update({
    "font.size": 11,
    "axes.labelsize": 11,
    "xtick.labelsize": 10,
    "ytick.labelsize": 10,
})


class ErroDados(Exception):
    """Problema que impede continuar. Mostrado ao utilizador sem traceback."""


# ---------------------------------------------------------------- utilitarios


def formatar_numero(valor) -> str:
    """Formata um numero a portuguesa: 1.234 e 1.234,56."""
    if valor is None:
        return "-"
    try:
        if pd.isna(valor):
            return "-"
    except (TypeError, ValueError):
        pass
    try:
        numero = float(valor)
    except (TypeError, ValueError):
        return str(valor)
    if abs(numero - round(numero)) < 1e-9:
        return f"{int(round(numero)):,}".replace(",", ".")
    texto = f"{numero:,.2f}"
    if texto.endswith("0"):
        texto = texto[:-1]  # 56,20% le-se mal; 56,2% le-se bem
    # troca os separadores anglo-saxonicos pelos portugueses
    return texto.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


def formatar_com_precisao(valor) -> str:
    """Formata guardando casas decimais suficientes para o numero dizer algo.

    Encontrado com dados reais de temperatura: um declive de 0,0005 arredondado
    a duas casas aparecia como «+0», e a previsao inteira colapsava em «entre 0
    e 1». O numero de casas passa a depender da grandeza.
    """
    numero = como_numero(valor)
    if numero is None:
        return formatar_numero(valor)

    magnitude = abs(numero)
    if abs(numero - round(numero)) < 1e-9:
        # inteiro: 1.214 le-se melhor do que 1.214,00
        casas = 0
    elif magnitude >= 1:
        # nunca menos de duas casas: cortar centimos fazia o numero impresso
        # contradizer o veredicto («0 abaixo da meta» com 0,40 de diferenca)
        casas = 2
    elif magnitude >= 0.01:
        casas = 4
    else:
        casas = 6

    texto = f"{numero:,.{casas}f}"
    if casas:
        # corta zeros a direita, mas nunca abaixo de duas casas: «0,4000» le-se
        # mal e «0,4» num valor monetario le-se pior
        inteiro, _, decimal = texto.partition(".")
        decimal = decimal.rstrip("0").ljust(2, "0")
        texto = f"{inteiro}.{decimal}"
    return texto.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


def formatar_categoria(valor) -> str:
    """Formata o rotulo de uma categoria; datas ficam em dd/mm/aaaa."""
    if isinstance(valor, pd.Timestamp):
        return valor.strftime("%d/%m/%Y")
    if valor is None:
        return "(vazio)"
    try:
        if pd.isna(valor):
            return "(vazio)"
    except (TypeError, ValueError):
        pass
    return str(valor)


def listar(itens) -> str:
    return ", ".join(f"«{item}»" for item in itens)


def onde_fica(folha: str, preposicao: str = "de") -> str:
    """«da folha X» num Excel, «do ficheiro» num CSV, que nao tem folhas.

    A preposicao muda com a frase: «a coluna X da folha Y», mas «não existe
    NA folha Y». Sem isto saia «não existe do ficheiro».
    """
    contracoes = {"de": ("do", "da"), "em": ("no", "na")}
    ficheiro, pasta = contracoes[preposicao]
    if folha == FonteCSV.NOME_UNICO:
        return f"{ficheiro} ficheiro"
    return f"{pasta} folha «{folha}»"


def listar_colunas(colunas) -> str:
    """Lista as colunas, mas nao despeja 93 nomes numa linha so."""
    nomes = [str(c) for c in colunas]
    if len(nomes) <= MAX_COLUNAS_LISTADAS:
        return listar(nomes)
    restantes = len(nomes) - MAX_COLUNAS_LISTADAS
    return f"{listar(nomes[:MAX_COLUNAS_LISTADAS])} e mais {restantes}"


def acrescentar(avisos: list[str], mensagem: str) -> None:
    """Junta um aviso so uma vez.

    Varios graficos leem a mesma folha; sem isto o mesmo aviso sairia
    repetido uma vez por grafico.
    """
    if mensagem not in avisos:
        avisos.append(mensagem)


def normalizar(texto) -> str:
    """Minusculas e sem acentos, para comparar rotulos."""
    sem_acentos = unicodedata.normalize("NFKD", str(texto))
    return "".join(c for c in sem_acentos if not unicodedata.combining(c)).strip().lower()


def limpar_numero(texto) -> str:
    """Tira simbolos de moeda e espacos, deixando so digitos e separadores."""
    limpo = str(texto).strip()
    for simbolo in SIMBOLOS_MOEDA:
        limpo = limpo.replace(simbolo, "")
    return limpo.replace(" ", "").replace(" ", "").strip()


def classificar_formato(limpo: str) -> str | None:
    """Diz que convencao de separadores um valor usa.

    Devolve 'pt' (1.250,00), 'en' (1,250.00), 'simples' (sem separadores),
    'ambiguo' (1.250 -- tanto pode ser mil duzentos e cinquenta como um
    virgula vinte e cinco) ou None se nem sequer parece um numero.
    """
    if not re.fullmatch(r"-?\d[\d.,]*", limpo):
        return None

    virgulas = limpo.count(",")
    pontos = limpo.count(".")

    if virgulas and pontos:
        # o ultimo separador a aparecer e o decimal
        return "pt" if limpo.rindex(",") > limpo.rindex(".") else "en"
    if virgulas:
        if virgulas > 1:
            return "en"  # 1,250,000 -- so pode ser separador de milhares
        return "ambiguo" if len(limpo.split(",")[-1]) == 3 else "pt"
    if pontos:
        if pontos > 1:
            return "pt"  # 1.250.000
        return "ambiguo" if len(limpo.split(".")[-1]) == 3 else "en"
    return "simples"


def aplicar_formato(limpo: str, convencao: str) -> float | None:
    if convencao == "pt":
        limpo = limpo.replace(".", "").replace(",", ".")
    elif convencao == "en":
        limpo = limpo.replace(",", "")
    try:
        return float(limpo)
    except ValueError:
        return None


def detetar_linha_cabecalho(fonte, folha: str) -> int:
    """Descobre em que linha comeca a tabela. Devolve o indice, base 0.

    Muita folha real tem o titulo do relatorio e a data de exportacao por
    cima da tabela. Sem isto, o titulo era lido como cabecalho e as colunas
    ficavam todas «Unnamed».

    O cabecalho e a PRIMEIRA linha, com pelo menos duas celulas e todas de
    texto, que esteja tao preenchida como a mais preenchida da amostra. Uma
    linha de titulo tem uma celula ou duas; o cabecalho tem a largura da
    tabela.

    A regra anterior comparava com a linha seguinte, e bastava a primeira
    linha de dados ter uma celula vazia para o cabecalho verdadeiro ser
    rejeitado -- encontrado com dados reais, onde isso e banal.

    Se nada servir, devolve 0: o comportamento de sempre.
    """
    try:
        amostra = fonte.ler(folha, cabecalho=None,
                            nrows=MAX_LINHAS_PROCURA_CABECALHO)
    except Exception:
        return 0
    if amostra.empty:
        return 0

    larguras = [len(amostra.iloc[i].dropna()) for i in range(len(amostra))]
    maxima = max(larguras, default=0)
    if maxima < 2:
        return 0

    for indice, largura in enumerate(larguras):
        if largura < maxima:
            continue
        preenchidas = amostra.iloc[indice].dropna()
        if all(isinstance(v, str) for v in preenchidas):
            return indice
    return 0




def converter_texto_formatado(coluna: pd.Series, onde: str,
                              nome: str) -> tuple[pd.Series, bool]:
    """Converte texto com separadores de milhares e moeda em numeros.

    A convencao e deduzida dos valores inequivocos da propria coluna. Se um
    «980,50 €» prova que a virgula e decimal, entao o «1.250» ao lado deixa
    de ser ambiguo. Se nada a prova, para com erro em vez de adivinhar: um
    palpite errado aqui nao da erro nenhum, da um relatorio mil vezes ao lado.
    """
    limpos = {}
    formatos = {}
    for indice, valor in coluna.items():
        if pd.isna(valor):
            continue
        limpo = limpar_numero(valor)
        limpos[indice] = limpo
        formatos[indice] = classificar_formato(limpo)

    convencoes = {f for f in formatos.values() if f in ("pt", "en")}
    ambiguos = [limpos[i] for i, f in formatos.items() if f == "ambiguo"]

    if len(convencoes) > 1:
        raise ErroDados(
            f"O {onde} usa «{nome}» como valor, mas a coluna mistura formatos de "
            "número incompatíveis (uns à portuguesa, outros à inglesa). "
            "Formata a coluna toda como número no Excel."
        )
    if ambiguos and not convencoes:
        exemplos = list(dict.fromkeys(ambiguos))[:3]
        raise ErroDados(
            f"O {onde} usa «{nome}» como valor, mas os números estão guardados "
            f"como texto num formato ambíguo: {listar(exemplos)}. «1.250» tanto "
            "pode ser mil duzentos e cinquenta como um vírgula vinte e cinco, e "
            "não vou adivinhar. Formata a coluna como número no Excel."
        )

    convencao = convencoes.pop() if convencoes else "simples"
    houve_formatacao = bool(convencoes) or convencao != "simples"

    valores = {}
    for indice, limpo in limpos.items():
        if formatos[indice] is None:
            valores[indice] = None
        else:
            valores[indice] = aplicar_formato(limpo, convencao)

    convertida = pd.Series(valores, dtype="float64").reindex(coluna.index)
    return convertida, houve_formatacao


def detetar_linha_de_total(serie: pd.Series, grafico: dict,
                           avisos: list[str]) -> None:
    """Avisa quando uma categoria parece ser a linha de totais da folha.

    Exportacoes de sistemas costumam deixar uma linha TOTAL no fim. Sem
    isto, essa linha vira uma categoria e o relatorio conta tudo duas
    vezes -- em silencio, que e o pior dos casos.

    Nao se apaga nada: avisa-se, e o aviso bloqueia a geracao ate haver
    uma decisao humana.
    """
    if len(serie) < 2:
        return

    suspeitas = [c for c in serie.index if normalizar(c) in ROTULOS_DE_TOTAL]

    # Sinal independente do nome: a ultima categoria vale a soma das outras.
    # So a partir de 5 categorias, e so na ultima. Com poucas categorias isto
    # acontece por acaso -- 10, 20 e 30 sao dados legitimos, e o 30 e a soma
    # dos outros dois. Linhas de totais ficam no fim da folha, nao no meio.
    if len(serie) >= 5:
        ultima = serie.index[-1]
        valor = float(serie.iloc[-1])
        resto = float(serie.sum()) - valor
        if resto > 0 and abs(valor - resto) <= abs(resto) * 1e-6:
            if ultima not in suspeitas:
                suspeitas.append(ultima)

    if suspeitas:
        acrescentar(avisos, (
            f"{grafico['titulo']}: {listar(suspeitas)} parece(m) ser linha(s) de "
            "totais da folha, e não categorias. Se ficarem, o gráfico conta os "
            "mesmos valores duas vezes. Apaga essa linha do Excel, ou confirma "
            "que é mesmo uma categoria."
        ))


# ------------------------------------------------------------------- leitura


def ler_plano(caminho: Path) -> dict:
    if not caminho.exists():
        raise ErroDados(f"O plano «{caminho}» não foi encontrado.")
    try:
        # utf-8-sig aceita o BOM que o Bloco de Notas e o PowerShell poem
        texto = caminho.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        raise ErroDados(
            f"O plano «{caminho}» não está em UTF-8. Grava-o outra vez em UTF-8."
        ) from None
    try:
        plano = json.loads(texto)
    except json.JSONDecodeError as erro:
        raise ErroDados(
            f"O plano «{caminho}» não é um JSON válido: {erro.msg} "
            f"(linha {erro.lineno}, coluna {erro.colno})."
        ) from None

    if not isinstance(plano, dict):
        raise ErroDados("O plano tem de ser um objeto JSON com chaves e valores.")
    titulo = plano.get("titulo_relatorio")
    if not isinstance(titulo, str) or not titulo.strip():
        raise ErroDados("Falta «titulo_relatorio» no plano, ou está vazio.")
    graficos = plano.get("graficos")
    if not isinstance(graficos, list) or not graficos:
        raise ErroDados("Falta «graficos» no plano, ou a lista está vazia.")

    analise = plano.get("analise", "completa")
    if analise not in ("completa", "curta", "detalhada"):
        raise ErroDados(
            f"O campo «analise» do plano é «{analise}», que não existe. "
            "Há «completa» (por omissão), «curta» e «detalhada» — esta última "
            "repete a análise inteira para cada série."
        )

    desconhecidas = sorted(set(plano) - {"titulo_relatorio", "graficos", "analise"})
    if desconhecidas:
        raise ErroDados(
            f"O plano tem campos que não existem: {listar(desconhecidas)}. "
            "Campos aceites: «titulo_relatorio», «graficos», «analise»."
        )

    for indice, grafico in enumerate(graficos, start=1):
        validar_grafico(grafico, indice)

    if analise == "curta":
        # sem isto, a meta e a previsao desapareciam sem uma palavra
        for indice, grafico in enumerate(graficos, start=1):
            pedidos = [c for c in ("meta", "previsao") if grafico.get(c)]
            if pedidos:
                raise ErroDados(
                    f"O gráfico {indice} pede {listar(pedidos)}, mas o plano tem "
                    "«analise»: «curta», que só escreve uma frase por gráfico. "
                    "Ou tiras esses campos, ou passas a análise a «completa» — "
                    "ignorá-los em silêncio seria pior."
                )
    return plano


def validar_grafico(grafico, indice: int) -> None:
    onde = f"gráfico {indice}"
    if not isinstance(grafico, dict):
        raise ErroDados(f"O {onde} não é um objeto JSON.")

    for chave in CHAVES_OBRIGATORIAS:
        if chave not in grafico:
            raise ErroDados(f"Falta «{chave}» no {onde}.")

    conhecidas = set(CHAVES_OBRIGATORIAS) | set(CHAVES_OPCIONAIS)
    desconhecidas = sorted(set(grafico) - conhecidas)
    if desconhecidas:
        raise ErroDados(
            f"O {onde} tem campos que não existem: {listar(desconhecidas)}. "
            f"Campos aceites: {listar(sorted(conhecidas))}."
        )

    if grafico["tipo"] not in TIPOS_GRAFICO:
        raise ErroDados(
            f"O {onde} pede o tipo «{grafico['tipo']}», que não existe. "
            f"Tipos disponíveis: {listar(TIPOS_GRAFICO)}."
        )
    if grafico["tipo"] in TIPOS_SEM_AGREGACAO:
        if grafico.get("agregacao"):
            raise ErroDados(
                f"O {onde} é de dispersão e traz «agregacao». Uma dispersão não "
                "agrupa nada: cruza duas colunas de números, um ponto por linha. "
                "Tira a «agregacao»."
            )
        if grafico.get("serie"):
            raise ErroDados(
                f"O {onde} é de dispersão e traz «serie». Ainda não há séries em "
                "gráficos de dispersão; faz um gráfico por série, ou usa «linhas»."
            )
        if not grafico.get("eixo_y"):
            raise ErroDados(
                f"O {onde} é de dispersão e precisa de «eixo_y»: são precisas duas "
                "colunas de números, uma para cada eixo."
            )
    elif not grafico.get("agregacao"):
        raise ErroDados(
            f"Falta «agregacao» no {onde}. "
            f"Agregações disponíveis: {listar(AGREGACOES)}."
        )
    elif grafico["agregacao"] not in AGREGACOES:
        raise ErroDados(
            f"O {onde} pede a agregação «{grafico['agregacao']}», que não existe. "
            f"Agregações disponíveis: {listar(AGREGACOES)}."
        )

    validar_filtro(grafico.get("filtro"), onde)
    if (grafico.get("agregacao") not in (None, "contagem")
            and not grafico.get("eixo_y")
            and not isinstance(grafico.get("serie"), list)):
        raise ErroDados(
            f"O {onde} usa a agregação «{grafico['agregacao']}» e por isso "
            "precisa de «eixo_y» (a coluna com os números), ou de «serie» como "
            "lista de colunas."
        )

    nota = grafico.get("nota")
    if nota is not None:
        if not isinstance(nota, str):
            raise ErroDados(f"A «nota» do {onde} tem de ser texto.")
        if any(caractere.isdigit() for caractere in nota):
            raise ErroDados(
                f"A «nota» do {onde} tem números. A nota serve para enquadrar "
                "(ex.: «período da campanha de verão»); os números do relatório "
                "são calculados a partir dos dados, para nunca desmentirem o gráfico."
            )

    if "tabela_dados" in grafico and not isinstance(grafico["tabela_dados"], bool):
        raise ErroDados(f"O campo «tabela_dados» do {onde} tem de ser true ou false.")

    if "eixo_temporal" in grafico and not isinstance(grafico["eixo_temporal"], bool):
        raise ErroDados(f"O campo «eixo_temporal» do {onde} tem de ser true ou false.")

    meta = grafico.get("meta")
    if meta is not None:
        if not isinstance(meta, dict):
            raise ErroDados(
                f"O campo «meta» do {onde} tem de ser um objeto com «valor» e "
                "«ambito», por exemplo {\"valor\": 700000, \"ambito\": \"total\"}."
            )
        desconhecidas = sorted(set(meta) - {"valor", "ambito"})
        if desconhecidas:
            raise ErroDados(
                f"A «meta» do {onde} tem campos que não existem: "
                f"{listar(desconhecidas)}. Só há «valor» e «ambito»."
            )
        alvo = como_numero(meta.get("valor"))
        if alvo is None:
            raise ErroDados(f"A «meta» do {onde} precisa de um «valor» numérico.")
        if alvo < 0:
            raise ErroDados(
                f"A «meta» do {onde} é {formatar_com_precisao(alvo)}. Uma meta "
                "negativa daria percentagens sem sentido (um total positivo "
                "ficaria com «-60% da meta»). Usa um valor de zero para cima."
            )
        ambito = meta.get("ambito", "total")
        if ambito == "serie" and not grafico.get("serie"):
            raise ErroDados(
                f"A «meta» do {onde} usa o âmbito «serie», mas o gráfico não tem "
                "séries. Usa «total» ou «categoria»."
            )
        if ambito == "categoria" and grafico.get("serie"):
            raise ErroDados(
                f"A «meta» do {onde} usa o âmbito «categoria» num gráfico com "
                "séries, e isso é ambíguo: «categoria» seriam os valores do eixo x "
                "ou as séries? Usa «serie» para comparar cada série com a meta, ou "
                "«total» para comparar o conjunto."
            )
        if ambito not in AMBITOS_META:
            raise ErroDados(
                f"A «meta» do {onde} pede o âmbito «{ambito}», que não existe. "
                f"Âmbitos disponíveis: {listar(AMBITOS_META)} — «total» compara "
                "o total do período, «categoria» compara cada categoria."
            )

    previsao = grafico.get("previsao")
    if previsao is not None:
        if not isinstance(previsao, int) or isinstance(previsao, bool) or previsao < 1:
            raise ErroDados(
                f"O campo «previsao» do {onde} tem de ser o número de períodos a "
                "prever (1, 2, 3…)."
            )

    coluna_serie = grafico.get("serie")
    if coluna_serie is not None:
        if isinstance(coluna_serie, list):
            if len(coluna_serie) < 2 or not all(
                    isinstance(c, str) and c.strip() for c in coluna_serie):
                raise ErroDados(
                    f"O campo «serie» do {onde} é uma lista de colunas e precisa "
                    "de pelo menos dois nomes de coluna."
                )
            if grafico.get("eixo_y"):
                raise ErroDados(
                    f"O {onde} dá «serie» como lista de colunas e também «eixo_y». "
                    "Com uma lista, cada coluna já é os valores de uma série — "
                    "tira o «eixo_y»."
                )
        elif not isinstance(coluna_serie, str) or not coluna_serie.strip():
            raise ErroDados(
                f"O campo «serie» do {onde} tem de ser o nome da coluna cujos valores "
                "dão as séries (ex.: «Canal»), ou uma lista de colunas quando cada "
                "categoria está na sua própria coluna."
            )

    periodo = grafico.get("coluna_periodo")
    if periodo is not None and (not isinstance(periodo, str) or not periodo.strip()):
        raise ErroDados(
            f"O campo «coluna_periodo» do {onde} tem de ser o nome da coluna com "
            "as datas ou os anos."
        )

    linha = grafico.get("linha_cabecalho")
    if linha is not None:
        if not isinstance(linha, int) or isinstance(linha, bool) or linha < 1:
            raise ErroDados(
                f"O campo «linha_cabecalho» do {onde} tem de ser o número da linha "
                "do Excel onde estão os nomes das colunas (1, 2, 3…)."
            )


class FonteExcel:
    """Le .xlsx e .xlsm. As macros de um .xlsm nunca sao executadas."""

    # o Excel guarda tipos; uma coluna de texto e mesmo uma escolha de quem gravou
    tudo_texto = False

    def __init__(self, caminho: Path):
        try:
            self._livro = pd.ExcelFile(caminho, engine="openpyxl")
        except PermissionError:
            raise ErroDados(
                f"Não consigo ler «{caminho}»: o ficheiro está aberto noutro programa "
                "(provavelmente o Excel). Fecha-o e tenta outra vez."
            ) from None
        except Exception as erro:  # ficheiro corrompido, zip invalido, etc.
            raise ErroDados(f"Não consigo abrir «{caminho}» como Excel: {erro}") from None
        self.notas: list[str] = []

    @property
    def folhas(self) -> list[str]:
        return list(self._livro.sheet_names)

    def ler(self, folha: str, cabecalho=0, nrows=None) -> pd.DataFrame:
        return self._livro.parse(folha, header=cabecalho, nrows=nrows)

    def bruto(self, folha: str, coluna: str, n_linhas: int,
              cabecalho: int = 0) -> list | None:
        """Valores tal como estao gravados na celula.

        E preciso porque o pandas decide sozinho que o texto «1.250» vale 1,25.
        Num Excel portugues isso quer dizer mil vezes menos, sem erro nenhum e
        sem aviso nenhum. A celula em bruto ainda sabe que aquilo era texto.

        Devolve None se nao conseguir alinhar com o que o pandas leu, para o
        resto do programa seguir pelo caminho normal em vez de arriscar.
        """
        primeira = cabecalho + 1  # openpyxl conta as linhas a partir de 1
        try:
            folha_bruta = self._livro.book[folha]
            cabecalhos = [
                c.value for c in next(folha_bruta.iter_rows(min_row=primeira,
                                                            max_row=primeira))
            ]
            indice = cabecalhos.index(coluna)
        except (AttributeError, KeyError, StopIteration, ValueError):
            return None

        valores = [
            linha[indice].value if indice < len(linha) else None
            for linha in folha_bruta.iter_rows(min_row=primeira + 1)
        ]
        # o pandas ignora linhas vazias no fim; so seguimos se bater certo
        while valores and valores[-1] is None and len(valores) > n_linhas:
            valores.pop()
        return valores if len(valores) == n_linhas else None


class FonteCSV:
    """Le CSV.

    Um CSV nao tem folhas, nao tem tipos e nao tem celulas: e tudo texto. Isso
    aqui e uma vantagem -- le-se tudo como texto e a conversao passa inteira
    pelo converter_texto_formatado, que so aceita o inequivoco. O caminho do
    CSV fica assim mais seguro do que o do Excel, nao menos.
    """

    NOME_UNICO = "(ficheiro)"
    # num CSV nao ha tipos: tudo o que la esta e texto, por definicao
    tudo_texto = True

    def __init__(self, caminho: Path):
        self.notas: list[str] = []
        texto, codificacao = self._ler_texto(caminho)
        separador = detetar_separador(texto)

        self.notas.append(
            f"CSV lido com «{separador}» como separador e codificação "
            f"{codificacao}."
        )
        # As linhas de metadados por cima da tabela tem menos colunas do que
        # ela, e o pandas recusa-se a ler um ficheiro assim. Conta-se primeiro
        # a linha mais larga e dao-se nomes a todas as colunas.
        try:
            largura = max(
                (len(linha) for linha in csv.reader(io.StringIO(texto),
                                                    delimiter=separador)),
                default=0,
            )
        except csv.Error:
            largura = 0

        try:
            self._tabela = pd.read_csv(
                io.StringIO(texto), sep=separador, header=None,
                names=range(largura) if largura else None,
                dtype=str,
                # so celulas mesmo vazias contam como vazias: com o default, a
                # categoria «NA» (North America) desaparecia do relatorio, e
                # «NULL» ou «nan» tambem. O que nao for numero e apanhado
                # depois pelo aviso de conversao, em vez de sumir em silencio.
                keep_default_na=False, na_values=[""],
                skip_blank_lines=False,
            )
        except Exception as erro:
            raise ErroDados(f"Não consigo interpretar «{caminho}» como CSV: {erro}") from None

        if self._tabela.empty:
            raise ErroDados(f"O ficheiro «{caminho}» não tem linhas nenhumas.")

    @staticmethod
    def _ler_texto(caminho: Path) -> tuple[str, str]:
        # utf-8-sig aceita o BOM; cp1252 e o que o Excel portugues costuma gravar
        for codificacao in ("utf-8-sig", "cp1252", "latin-1"):
            try:
                return caminho.read_text(encoding=codificacao), codificacao
            except UnicodeDecodeError:
                continue
            except PermissionError:
                raise ErroDados(
                    f"Não consigo ler «{caminho}»: o ficheiro está aberto noutro "
                    "programa. Fecha-o e tenta outra vez."
                ) from None
        raise ErroDados(
            f"Não consigo descodificar «{caminho}». Tentei UTF-8, cp1252 e latin-1."
        )

    @property
    def folhas(self) -> list[str]:
        return [self.NOME_UNICO]

    def ler(self, folha: str, cabecalho=0, nrows=None) -> pd.DataFrame:
        tabela = self._tabela if nrows is None else self._tabela.head(nrows)
        if cabecalho is None:
            return tabela.reset_index(drop=True)
        corpo = tabela.iloc[cabecalho + 1:].reset_index(drop=True)
        corpo.columns = nomes_unicos([
            str(c) if pd.notna(c) else f"Unnamed: {i}"
            for i, c in enumerate(tabela.iloc[cabecalho])
        ])
        return corpo

    def bruto(self, folha: str, coluna: str, n_linhas: int,
              cabecalho: int = 0) -> list | None:
        corpo = self.ler(folha, cabecalho)
        if coluna not in corpo.columns or len(corpo) != n_linhas:
            return None
        return [None if pd.isna(v) else v for v in corpo[coluna]]


def nomes_unicos(nomes: list[str]) -> list[str]:
    """Desambigua cabecalhos repetidos, como o pandas ja faz no Excel.

    Exportacoes de plataformas trazem colunas com o mesmo nome a torto e a
    direito. Sem isto, df[coluna] devolvia um DataFrame em vez de uma coluna e
    o programa rebentava com um traceback.
    """
    vistos: dict[str, int] = {}
    saida = []
    for nome in nomes:
        if nome in vistos:
            vistos[nome] += 1
            saida.append(f"{nome}.{vistos[nome]}")
        else:
            vistos[nome] = 0
            saida.append(nome)
    return saida


def detetar_separador(texto: str) -> str:
    """Descobre o separador do CSV. O Excel portugues exporta com «;»."""
    amostra = "\n".join(texto.splitlines()[:20])
    try:
        return csv.Sniffer().sniff(amostra, delimiters=";,\t|").delimiter
    except csv.Error:
        pass
    # recurso: conta ocorrencias fora de aspas e fica com o mais frequente
    contagens = {}
    for candidato in (";", ",", "\t", "|"):
        total, dentro_de_aspas = 0, False
        for caractere in amostra:
            if caractere == '"':
                dentro_de_aspas = not dentro_de_aspas
            elif caractere == candidato and not dentro_de_aspas:
                total += 1
        contagens[candidato] = total
    melhor = max(contagens, key=contagens.get)
    return melhor if contagens[melhor] > 0 else ","


def validar_filtro(filtro, onde: str) -> None:
    if filtro is None:
        return
    if not isinstance(filtro, dict):
        raise ErroDados(
            f"O «filtro» do {onde} tem de ser um objeto, por exemplo "
            '{"coluna": "Ano", "igual_a": 2025}.'
        )
    desconhecidas = sorted(set(filtro) - {"coluna", "igual_a", "de", "ate"})
    if desconhecidas:
        raise ErroDados(
            f"O «filtro» do {onde} tem campos que não existem: "
            f"{listar(desconhecidas)}. Só há «coluna», «igual_a», «de» e «ate»."
        )
    if not isinstance(filtro.get("coluna"), str) or not filtro["coluna"].strip():
        raise ErroDados(f"O «filtro» do {onde} precisa de «coluna».")
    if not any(chave in filtro for chave in ("igual_a", "de", "ate")):
        raise ErroDados(
            f"O «filtro» do {onde} não diz o que filtrar. Usa «igual_a» (um valor "
            "ou uma lista), «de» e «ate» (intervalo), ou os dois."
        )


def aplicar_filtro(df: pd.DataFrame, grafico: dict, onde: str, folha: str,
                   notas: list[str]) -> pd.DataFrame:
    """Filtra as linhas, e escreve SEMPRE no relatorio o que ficou de fora.

    Filtrar e esconder dados. Um relatorio que mostra 2025 sem dizer que
    ignorou 2023 e 2024 e enganador, por isso o filtro fica registado.
    """
    filtro = grafico.get("filtro")
    if not filtro:
        return df

    coluna = filtro["coluna"]
    if coluna not in df.columns:
        raise ErroDados(
            f"O filtro do {onde} usa a coluna «{coluna}», que não existe "
            f"{onde_fica(folha, 'em')}. "
            f"Colunas disponíveis: {listar_colunas(df.columns)}."
        )

    antes = len(df)
    condicoes = []
    valores = df[coluna]

    if "igual_a" in filtro:
        alvo = filtro["igual_a"]
        alvos = alvo if isinstance(alvo, list) else [alvo]
        # compara como texto: no CSV vem tudo como texto e 2025 != "2025"
        df = df[valores.map(lambda v: str(v) in {str(a) for a in alvos})]
        condicoes.append(f"«{coluna}» igual a {listar([str(a) for a in alvos])}")

    for chave, simbolo, comparar in (("de", "≥", lambda s, l: s >= l),
                                     ("ate", "≤", lambda s, l: s <= l)):
        if chave in filtro:
            numeros = pd.to_numeric(df[coluna], errors="coerce")
            limite = como_numero(filtro[chave])
            if limite is None or numeros.notna().sum() == 0:
                raise ErroDados(
                    f"O filtro do {onde} compara «{coluna}» com "
                    f"«{filtro[chave]}», mas isso exige números dos dois lados."
                )
            df = df[comparar(numeros, limite)]
            condicoes.append(f"«{coluna}» {simbolo} {formatar_com_precisao(limite)}")

    if df.empty:
        raise ErroDados(
            f"O filtro do {onde} ({'; '.join(condicoes)}) não deixou nenhuma linha."
        )

    notas.append(
        f"{grafico['titulo']}: filtro aplicado — {'; '.join(condicoes)}. "
        f"Usadas {formatar_numero(len(df))} de {formatar_numero(antes)} linhas; "
        f"{formatar_numero(antes - len(df))} ficaram de fora."
    )
    return df


def abrir_dados(caminho: Path):
    if not caminho.exists():
        raise ErroDados(f"O ficheiro «{caminho}» não foi encontrado.")

    extensao = caminho.suffix.lower()
    if extensao in (".xlsx", ".xlsm"):
        return FonteExcel(caminho)
    if extensao in (".csv", ".txt", ".tsv"):
        return FonteCSV(caminho)

    raise ErroDados(
        f"Não sei ler ficheiros «{extensao or 'sem extensão'}». "
        "Suportados: .xlsx, .xlsm e .csv. "
        + ("O .xls é o formato antigo do Excel; abre-o e grava como .xlsx."
           if extensao == ".xls" else "Grava o ficheiro num destes formatos.")
    )


# --------------------------------------------------------------- preparacao


def escolher_folha(fonte, grafico: dict, onde: str, notas: list[str]) -> str:
    """Decide de que folha se le. Um CSV nao tem folhas; um Excel pode ter uma so."""
    disponiveis = fonte.folhas
    pedida = grafico.get("folha")

    if fonte.tudo_texto:
        if pedida:
            notas.append(
                f"{grafico['titulo']}: o plano indica a folha «{pedida}», mas um CSV "
                "não tem folhas. Foi ignorada."
            )
        return disponiveis[0]

    if pedida is None:
        if len(disponiveis) == 1:
            return disponiveis[0]
        raise ErroDados(
            f"O {onde} não diz de que folha ler, e o ficheiro tem "
            f"{len(disponiveis)}: {listar(disponiveis)}."
        )

    if pedida not in disponiveis:
        raise ErroDados(
            f"O {onde} pede a folha «{pedida}», que não existe no ficheiro. "
            f"Folhas disponíveis: {listar(disponiveis)}."
        )
    return pedida


def reconhecer_datas(df: pd.DataFrame, coluna: str, onde: str,
                     notas: list[str]) -> pd.DataFrame:
    """Converte o eixo x em datas quando for texto de datas com convencao provada."""
    if pd.api.types.is_datetime64_any_dtype(df[coluna]):
        return df
    if df[coluna].isna().any() or not len(df):
        return df

    amostra = list(df[coluna])
    if not all(partes_de_data(v) is not None for v in amostra):
        return df

    convertidas = converter_datas_texto(amostra, onde, coluna)
    if convertidas is None:
        notas.append(
            f"A coluna «{coluna}» tem datas em texto, mas nenhum dia passa de 12: "
            "não é possível saber se é dia/mês ou mês/dia. Ficou como categoria, "
            "sem análise de tendência — trocar Janeiro por Fevereiro daria um "
            "resultado errado com ar de certo."
        )
        return df

    df = df.copy()
    df[coluna] = convertidas
    acrescentar(notas, f"A coluna «{coluna}» foi reconhecida como datas.")
    return df


def agregar(df: pd.DataFrame, eixo_x: str, eixo_y: str | None,
            agregacao: str) -> pd.Series:
    """Agrupa por eixo_x pela ordem de aparicao no ficheiro."""
    if eixo_y is None:
        return df.groupby(eixo_x, sort=False).size()
    grupo = df.groupby(eixo_x, sort=False)[eixo_y]
    return {"soma": grupo.sum, "media": grupo.mean,
            "contagem": grupo.count}[agregacao]()


def series_de_colunas(df: pd.DataFrame, grafico: dict, eixo_x: str,
                      agregacao: str, folha: str, onde: str,
                      avisos: list[str]) -> dict:
    """Series a partir de VARIAS COLUNAS, para dados em formato largo.

    Muita folha real tem uma coluna por regiao, por mes ou por canal, em vez
    de uma coluna com o nome da categoria. Encontrado com dados publicos
    reais: sem isto, «comparar por regiao» era simplesmente impossivel.
    """
    colunas = grafico["serie"]
    em_falta = [c for c in colunas if c not in df.columns]
    if em_falta:
        raise ErroDados(
            f"O {onde} pede as colunas {listar(em_falta)} como séries, e não "
            f"existem {onde_fica(folha, 'em')}. "
            f"Colunas disponíveis: {listar_colunas(df.columns)}."
        )
    if eixo_x in colunas:
        raise ErroDados(
            f"O {onde} usa «{eixo_x}» como eixo x e também como série. "
            "Uma coluna não pode ser as duas coisas."
        )

    series = {}
    for nome in colunas:
        limpo = df.dropna(subset=[nome])
        if limpo.empty:
            acrescentar(avisos, (
                f"A coluna «{nome}» não tem valor nenhum preenchido; "
                "essa série ficou de fora."
            ))
            continue
        agregada = agregar(limpo, eixo_x, nome, agregacao)
        if not agregada.empty:
            series[str(nome)] = agregada

    if len(series) < 2:
        raise ErroDados(
            f"O {onde} pede séries a partir de colunas, mas sobrou "
            f"{len(series)}. São precisas pelo menos duas para haver comparação."
        )
    if len(series) > MAX_SERIES:
        acrescentar(avisos, (
            f"{grafico['titulo']}: são {len(series)} séries no mesmo gráfico. "
            f"Acima de {MAX_SERIES} as linhas confundem-se umas com as outras."
        ))
    return series


def detetar_acumulado(serie: pd.Series, grafico: dict, agregacao: str,
                      avisos: list[str]) -> None:
    """Avisa quando a serie parece um acumulado e se esta a soma-la.

    Encontrado com dados publicos reais de casos acumulados: o relatorio dizia
    «o total é 720.849.837» quando o total verdadeiro era 5.372.111. Somar uma
    coluna acumulada da um numero sem significado nenhum -- e nao ha erro
    nenhum a denuncia-lo, so um numero absurdo com ar de conta certa.

    O sinal: uma serie longa que nunca desce.
    """
    if agregacao != "soma" or len(serie) < MIN_PERIODOS_ACUMULADO:
        return
    valores = [float(v) for v in serie.values]
    if any(b < a for a, b in zip(valores, valores[1:])):
        return
    if valores[-1] <= valores[0]:
        return

    coluna = grafico.get("eixo_y")
    qual = f"a coluna «{coluna}»" if coluna and coluna != "_conjunto" else "esta série"
    acrescentar(avisos, (
        f"{grafico['titulo']}: {qual} nunca desce ao "
        f"longo de {formatar_numero(len(serie))} períodos — parece um valor "
        "acumulado. Somar um acumulado dá um número sem significado "
        f"({formatar_com_precisao(sum(valores))} em vez de "
        f"{formatar_com_precisao(valores[-1])}). Se for esse o caso, usa o valor "
        "do último período, ou uma coluna com os valores de cada período."
    ))


def separar_series(df: pd.DataFrame, grafico: dict, eixo_x: str,
                   eixo_y: str | None, agregacao: str, folha: str, onde: str,
                   avisos: list[str], notas: list[str]) -> dict:
    """Devolve {nome_da_serie: serie_agregada}.

    Sem o campo «serie» devolve uma entrada so, com nome vazio -- e assim todo
    o codigo de analise, que trabalha sobre uma pd.Series, continua igual e
    apenas passa a ser chamado uma vez por serie.
    """
    coluna = grafico.get("serie")
    if not coluna:
        return {"": agregar(df, eixo_x, eixo_y, agregacao)}

    if isinstance(coluna, list):
        return series_de_colunas(df, grafico, eixo_x, agregacao, folha, onde, avisos)

    if coluna not in df.columns:
        raise ErroDados(
            f"O {onde} pede «{coluna}» como coluna das séries, que não existe "
            f"{onde_fica(folha, 'em')}. Colunas disponíveis: {listar_colunas(df.columns)}."
        )
    if coluna == eixo_x:
        raise ErroDados(
            f"O {onde} usa «{coluna}» como eixo x e como série ao mesmo tempo. "
            "Cada série ficaria com um ponto só."
        )
    if grafico["tipo"] == "circular":
        raise ErroDados(
            f"O {onde} é circular e pede séries. Um gráfico circular mostra a "
            "repartição de um total, não a evolução de várias séries. Usa "
            "«linhas» ou «barras», ou tira o campo «serie»."
        )

    limpo = df.dropna(subset=[coluna])
    perdidas = len(df) - len(limpo)
    if perdidas:
        acrescentar(avisos, (
            f"{perdidas} linha(s) foram ignoradas por não terem valor em "
            f"«{coluna}», a coluna das séries."
        ))
    if limpo.empty:
        raise ErroDados(f"O {onde} ficou sem dados depois de agrupar por «{coluna}».")

    series = {}
    for nome in limpo[coluna].drop_duplicates():  # ordem de aparicao no ficheiro
        parte = limpo[limpo[coluna] == nome]
        agregada = agregar(parte, eixo_x, eixo_y, agregacao)
        if not agregada.empty:
            series[formatar_categoria(nome)] = agregada

    if not series:
        raise ErroDados(f"O {onde} não produziu série nenhuma a partir de «{coluna}».")

    if len(series) > MAX_SERIES:
        acrescentar(avisos, (
            f"{grafico['titulo']}: «{coluna}» daria {len(series)} séries no mesmo "
            f"gráfico. Acima de {MAX_SERIES} as linhas confundem-se umas com as "
            "outras. Considera filtrar, ou agrupar as mais pequenas."
        ))

    curtas = [n for n, s in series.items() if len(s) < 2]
    if curtas:
        notas.append(
            f"{grafico['titulo']}: {listar(curtas)} tem um ponto só; "
            "a tendência dessas séries não é calculada."
        )
    return series


def preparar_dispersao(df: pd.DataFrame, grafico: dict, eixo_x: str,
                       eixo_y: str, folha: str, onde: str, fonte,
                       avisos: list[str], notas: list[str]) -> tuple:
    """Prepara um grafico de dispersao: um ponto por linha, sem agrupar.

    As duas colunas tem de ser numeros. Nao ha categorias nem agregacao, por
    isso tambem nao ha analise temporal, sazonalidade nem previsao.
    """
    if eixo_y not in df.columns:
        raise ErroDados(
            f"O {onde} pede a coluna «{eixo_y}», que não existe "
            f"{onde_fica(folha, 'em')}. "
            f"Colunas disponíveis: {listar_colunas(df.columns)}."
        )

    for coluna in (eixo_x, eixo_y):
        df = converter_para_numero(df, coluna, folha, onde, avisos, notas,
                                   tudo_texto=fonte.tudo_texto)

    antes = len(df)
    df = df.dropna(subset=[eixo_x, eixo_y])
    if antes - len(df):
        acrescentar(avisos, (
            f"{antes - len(df)} linha(s) {onde_fica(folha)} foram ignoradas por "
            f"terem células vazias em {listar([eixo_x, eixo_y])}."
        ))
    if len(df) < 3:
        raise ErroDados(
            f"O {onde} é de dispersão e ficou com {len(df)} ponto(s). "
            "São precisos pelo menos 3 para o gráfico dizer alguma coisa."
        )

    serie = pd.Series(df[eixo_y].to_numpy(), index=df[eixo_x].to_numpy())
    return {"": serie}, len(df), False, None, serie


def preparar_dados(fonte, grafico: dict, indice: int,
                   avisos: list[str], notas: list[str]) -> tuple[pd.Series, int]:
    """Le a folha, valida colunas e devolve a serie agregada e o nº de linhas usadas.

    A ordem das categorias e a ordem de aparicao no ficheiro (sort=False).
    Ordenar por ordem alfabetica estragaria series temporais: Abril viria
    antes de Janeiro.
    """
    onde = f"gráfico {indice} («{grafico['titulo']}»)"
    folha = escolher_folha(fonte, grafico, onde, notas)

    pedida = grafico.get("linha_cabecalho")
    if pedida is None:
        linha_cabecalho = detetar_linha_cabecalho(fonte, folha)
        if linha_cabecalho > 0:
            acrescentar(avisos, (
                f"A tabela {onde_fica(folha)} não começa na primeira linha. "
                f"Usei a linha {linha_cabecalho + 1} como cabeçalho, e ignorei as "
                f"{linha_cabecalho} de cima (costumam ser o título e a data de "
                "exportação). Confirma que é essa a linha certa."
            ))
    else:
        linha_cabecalho = pedida - 1  # o plano conta as linhas como o Excel

    df = fonte.ler(folha, cabecalho=linha_cabecalho)
    if df.empty:
        raise ErroDados(f"A tabela {onde_fica(folha)} não tem linhas nenhumas.")

    fantasma = [c for c in df.columns if str(c).startswith("Unnamed:")]
    if fantasma:
        acrescentar(avisos, (
            f"A tabela {onde_fica(folha)} tem {len(fantasma)} coluna(s) sem cabeçalho "
            f"({listar(fantasma)}). Costuma ser sinal de células soltas ao lado "
            "da tabela. Foram ignoradas."
        ))

    df = aplicar_filtro(df, grafico, onde, folha, notas)

    eixo_x = grafico["eixo_x"]
    if eixo_x not in df.columns:
        raise ErroDados(
            f"O {onde} pede a coluna «{eixo_x}», que não existe {onde_fica(folha, 'em')}. "
            f"Colunas disponíveis: {listar_colunas(df.columns)}."
        )

    df = reconhecer_datas(df, eixo_x, onde, notas)

    agregacao = grafico.get("agregacao")
    eixo_y = grafico.get("eixo_y")

    if grafico["tipo"] in TIPOS_SEM_AGREGACAO:
        return preparar_dispersao(df, grafico, eixo_x, eixo_y, folha, onde,
                                  fonte, avisos, notas)

    if isinstance(grafico.get("serie"), list):
        # formato largo: cada coluna da lista e uma serie. Converte-se cada uma,
        # e o conjunto passa a ser a combinacao delas linha a linha.
        for coluna in grafico["serie"]:
            if coluna not in df.columns:
                raise ErroDados(
                    f"O {onde} pede a coluna «{coluna}» como série, que não existe "
                    f"{onde_fica(folha, 'em')}. "
                    f"Colunas disponíveis: {listar_colunas(df.columns)}."
                )
            df = converter_para_numero(df, coluna, folha, onde, avisos, notas,
                                       tudo_texto=fonte.tudo_texto)
        eixo_y = "_conjunto"
        combinar = df[grafico["serie"]]
        df = df.copy()
        df[eixo_y] = combinar.mean(axis=1) if agregacao == "media" else combinar.sum(axis=1)
        colunas_necessarias = [eixo_x]
    elif eixo_y is not None:
        if eixo_y not in df.columns:
            raise ErroDados(
                f"O {onde} pede a coluna «{eixo_y}», que não existe {onde_fica(folha, 'em')}. "
                f"Colunas disponíveis: {listar_colunas(df.columns)}."
            )
        if not fonte.tudo_texto:
            # so faz sentido no Excel: e la que o pandas decide sozinho que o
            # texto «1.250» vale 1,25. No CSV os valores ja chegam como texto e
            # passam todos pelo converter_para_numero a seguir.
            df = proteger_de_texto_reinterpretado(
                df, fonte, folha, eixo_y, onde, avisos, notas, linha_cabecalho)
        df = converter_para_numero(df, eixo_y, folha, onde, avisos, notas,
                                   tudo_texto=fonte.tudo_texto)
        colunas_necessarias = [eixo_x, eixo_y]
    else:
        colunas_necessarias = [eixo_x]

    antes = len(df)
    df = df.dropna(subset=colunas_necessarias)
    ignoradas = antes - len(df)
    if ignoradas:
        acrescentar(avisos, (
            f"{ignoradas} linha(s) {onde_fica(folha)} foram ignoradas por terem "
            f"células vazias em {listar(colunas_necessarias)}."
        ))
    if df.empty:
        raise ErroDados(
            f"Depois de ignorar as linhas com células vazias, o {onde} ficou sem dados."
        )

    serie = agregar(df, eixo_x, eixo_y, agregacao)
    if serie.empty:
        raise ErroDados(f"O {onde} não produziu nenhuma categoria depois de agrupar.")

    # O «conjunto» e sempre esta agregacao, feita sobre os dados em bruto e a
    # ignorar a coluna das series. Nao e a soma das series: com agregacao
    # «media», somar as medias de cada serie daria um numero errado.
    series = separar_series(df, grafico, eixo_x, eixo_y, agregacao, folha,
                            onde, avisos, notas)

    detetar_linha_de_total(serie, grafico, avisos)
    detetar_acumulado(serie, grafico, agregacao, avisos)

    temporal = grafico.get("eixo_temporal")
    if temporal is None:
        temporal = parece_temporal(serie.index)
    if temporal:
        notas.append(f"{grafico['titulo']}: «{eixo_x}» tratado como linha do tempo.")

    serie_anual = calcular_serie_anual(df, grafico, eixo_x, eixo_y, folha, avisos)

    if len(df) > len(serie):
        notas.append(
            f"{grafico['titulo']}: {len(df)} linhas agrupadas em {len(serie)} "
            f"categorias de «{eixo_x}» ({agregacao})."
        )

    # so em barras: numa linha, 64 pontos leem-se bem; 64 barras nao
    if grafico["tipo"] == "barras" and len(serie) > MAX_CATEGORIAS_LEGIVEIS:
        acrescentar(avisos, (
            f"{grafico['titulo']}: o gráfico ficaria com {len(serie)} categorias de "
            f"«{eixo_x}». Acima de {MAX_CATEGORIAS_LEGIVEIS} as barras ficam finas "
            "como cabelos e os rótulos ilegíveis. Considera agrupar, ou usar uma "
            "coluna com menos categorias."
        ))

    if grafico.get("tabela_dados") and len(serie) > MAX_LINHAS_TABELA:
        acrescentar(avisos, (
            f"{grafico['titulo']}: a tabela de dados mostraria as primeiras "
            f"{MAX_LINHAS_TABELA} de {len(serie)} categorias; "
            f"{len(serie) - MAX_LINHAS_TABELA} ficariam de fora."
        ))

    if grafico["tipo"] == "circular":
        if (serie < 0).any():
            negativas = [formatar_categoria(c) for c in serie[serie < 0].index]
            raise ErroDados(
                f"O {onde} é circular, mas há valores negativos em {listar(negativas)}. "
                "Uma fatia não pode ser negativa. Usa barras, ou filtra esses valores."
            )
        if (serie == 0).all():
            raise ErroDados(f"O {onde} é circular e todos os valores são zero.")
        if len(serie) > MAX_FATIAS_CIRCULAR:
            acrescentar(avisos, (
                f"{grafico['titulo']}: o gráfico circular ficaria com {len(serie)} "
                f"fatias. Acima de {MAX_FATIAS_CIRCULAR} deixa de se ler. "
                "Considera barras, ou agrupar as categorias mais pequenas."
            ))

    return series, len(df), bool(temporal), serie_anual, serie


def calcular_serie_anual(df: pd.DataFrame, grafico: dict, eixo_x: str,
                         eixo_y: str | None, folha: str,
                         avisos: list[str]) -> pd.Series | None:
    """Agrega por ano, para a analise ano a ano.

    A coluna dos anos e a indicada em «coluna_periodo»; se nao houver, tenta o
    proprio eixo_x. Devolve None quando nao ha anos para tirar dali.
    """
    coluna = grafico.get("coluna_periodo") or eixo_x
    if coluna not in df.columns:
        raise ErroDados(
            f"O gráfico «{grafico['titulo']}» pede «{coluna}» como coluna de "
            f"períodos, mas essa coluna não existe {onde_fica(folha, 'em')}. "
            f"Colunas disponíveis: {listar_colunas(df.columns)}."
        )

    anos = df[coluna].map(extrair_ano)
    if anos.isna().all():
        if grafico.get("coluna_periodo"):
            acrescentar(avisos, (
                f"Não consegui tirar anos da coluna «{coluna}» {onde_fica(folha)}. "
                "A análise ano a ano ficou de fora deste gráfico."
            ))
        return None

    trabalho = df.assign(_ano=anos).dropna(subset=["_ano"])
    if eixo_y is None:
        return trabalho.groupby("_ano", sort=True).size()

    grupo = trabalho.groupby("_ano", sort=True)[eixo_y]
    agregacao = grafico["agregacao"]
    return {"soma": grupo.sum, "media": grupo.mean, "contagem": grupo.count}[agregacao]()


def proteger_de_texto_reinterpretado(df: pd.DataFrame, fonte,
                                     folha: str, coluna: str, onde: str,
                                     avisos: list[str], notas: list[str],
                                     linha_cabecalho: int = 0) -> pd.DataFrame:
    """Reconverte a coluna a partir das celulas em bruto quando la havia texto.

    So para ficheiros Excel: e la que o pandas decide sozinho que o texto
    «1.250» vale 1,25. So faz alguma coisa se a coluna tiver mesmo celulas de
    texto com separadores ou simbolos de moeda; numeros a serio nao passam por
    aqui.
    """
    brutos = fonte.bruto(folha, coluna, len(df), linha_cabecalho)
    if brutos is None:
        return df

    formatados = [
        v for v in brutos
        if isinstance(v, str) and classificar_formato(limpar_numero(v)) in
        ("pt", "en", "ambiguo")
    ]
    if not formatados:
        return df

    coluna_bruta = pd.Series(brutos, index=df.index, dtype="object")
    convertida, _ = converter_texto_formatado(coluna_bruta, onde, coluna)

    mensagem = (
        f"A coluna «{coluna}» {onde_fica(folha)} tem números gravados como texto "
        "(com separadores de milhares ou símbolos de moeda). Foram convertidos a "
        "partir da célula original."
    )
    alterou_significado = False
    if pd.api.types.is_numeric_dtype(df[coluna]):
        antes = pd.to_numeric(df[coluna], errors="coerce")
        diferentes = int((~antes.sub(convertida).abs().le(1e-9)).sum())
        if diferentes:
            alterou_significado = True
            mensagem += (
                f" Atenção: {diferentes} valor(es) tinham sido lidos com o ponto "
                "como separador decimal, o que dava mil vezes menos. Confirma-os "
                "no relatório."
            )

    acrescentar(avisos, mensagem)

    df = df.copy()
    df[coluna] = convertida
    return df


def converter_para_numero(df: pd.DataFrame, coluna: str, folha: str,
                          onde: str, avisos: list[str], notas: list[str],
                          tudo_texto: bool = False) -> pd.DataFrame:
    """Garante que a coluna e numerica. Nunca soma texto em silencio.

    Num CSV e tudo texto por definicao, e o aviso «estava guardada como texto»
    dispararia em todas as colunas de todos os ficheiros. Um aviso que aparece
    sempre deixa de significar alguma coisa, e o bloqueio virava ruido. Nesse
    caso a conversao vai para as notas; a perda de celulas continua a ser aviso.
    """
    if pd.api.types.is_numeric_dtype(df[coluna]):
        return df

    if pd.api.types.is_datetime64_any_dtype(df[coluna]):
        raise ErroDados(
            f"O {onde} usa «{coluna}» como valor, mas essa coluna tem datas, "
            "não números."
        )

    convertida, com_formatacao = converter_texto_formatado(df[coluna], onde, coluna)
    validos = int(convertida.notna().sum())
    if validos == 0:
        exemplos = [str(v) for v in df[coluna].dropna().unique()[:3]]
        raise ErroDados(
            f"O {onde} usa «{coluna}» como valor, mas essa coluna não tem números. "
            f"Exemplos do que lá está: {listar(exemplos)}."
        )

    ja_vazias = int(df[coluna].isna().sum())
    perdidas = int(convertida.isna().sum()) - ja_vazias
    mensagem = (
        f"A coluna «{coluna}» {onde_fica(folha)} estava guardada como texto e "
        "foi convertida para número"
        + (" (separadores de milhares e símbolos de moeda incluídos)." if com_formatacao
           else ".")
    )
    if perdidas > 0:
        nao_numericas = df.loc[convertida.isna() & df[coluna].notna(), coluna]
        exemplos = [str(v) for v in nao_numericas.unique()[:3]]
        mensagem += (
            f" {perdidas} célula(s) não eram números e foram ignoradas "
            f"(exemplos: {listar(exemplos)})."
        )

    if tudo_texto and perdidas == 0:
        # num CSV isto e o normal, nao uma surpresa
        acrescentar(notas, mensagem)
    else:
        acrescentar(avisos, mensagem)

    df = df.copy()
    df[coluna] = convertida
    return df


# ----------------------------------------------------------------- desenho


def desenhar(serie: pd.Series, series: dict, grafico: dict, destino: Path) -> None:
    """Desenha o grafico em PNG. O titulo nao entra na imagem: vai no Word.

    «series» e o dicionario {nome: pd.Series}. Com uma entrada de nome vazio
    desenha como sempre desenhou; com varias, uma linha ou um grupo de barras
    por serie, com legenda.
    """
    nomes = [n for n in series if n != ""]
    rotulos = [formatar_categoria(c) for c in serie.index]
    valores = list(serie.values)
    tipo = grafico["tipo"]

    # a figura entra no Word com LARGURA_IMAGEM; desenhar muito maior do que
    # isso encolhe o texto ate ficar ilegivel na pagina impressa
    figura, eixos = plt.subplots(figsize=(7, 4), dpi=200)

    if tipo == "dispersao":
        eixos.scatter([como_numero(c) for c in serie.index], valores,
                      color=CORES[0], alpha=0.7, edgecolors="white", linewidths=0.5)
        eixos.grid(linestyle=":", alpha=0.6)
    elif nomes and tipo in ("barras", "barras_horizontais", "linhas", "area"):
        desenhar_series(eixos, series, serie.index, rotulos, tipo)
    elif tipo == "barras":
        eixos.bar(rotulos, valores, color=CORES[0])
        eixos.grid(axis="y", linestyle=":", alpha=0.6)
    elif tipo == "barras_horizontais":
        # deitadas de propositio: nomes compridos nao cabem numa barra em pe
        eixos.barh(rotulos, valores, color=CORES[0])
        eixos.invert_yaxis()  # o primeiro do ficheiro fica em cima
        eixos.grid(axis="x", linestyle=":", alpha=0.6)
    elif tipo == "linhas":
        eixos.plot(rotulos, valores, marker="o", color=CORES[0], linewidth=2)
        eixos.grid(linestyle=":", alpha=0.6)
    elif tipo == "area":
        eixos.fill_between(range(len(valores)), valores, color=CORES[0], alpha=0.35)
        eixos.plot(range(len(valores)), valores, color=CORES[0], linewidth=2)
        eixos.set_xticks(range(len(rotulos)))
        eixos.set_xticklabels(rotulos)
        eixos.grid(linestyle=":", alpha=0.6)
    else:  # circular
        eixos.pie(
            valores,
            labels=rotulos,
            # percentagem a portuguesa, com virgula, como o resto do relatorio
            autopct=lambda p: f"{formatar_numero(round(p, 1))}%",
            startangle=90,
            counterclock=False,
        )
        eixos.axis("equal")

    if tipo == "barras_horizontais":
        eixos.set_axisbelow(True)
        eixos.set_ylabel(str(grafico["eixo_x"]))
        eixos.set_xlabel(rotulo_do_eixo_y(grafico))
        eixos.xaxis.set_major_formatter(
            matplotlib.ticker.FuncFormatter(lambda v, _: formatar_numero(v))
        )
    elif tipo != "circular":
        eixos.set_axisbelow(True)
        eixos.set_xlabel(str(grafico["eixo_x"]))
        eixos.set_ylabel(rotulo_do_eixo_y(grafico))
        eixos.yaxis.set_major_formatter(
            matplotlib.ticker.FuncFormatter(lambda v, _: formatar_numero(v))
        )
        if tipo != "dispersao":
            rodar_rotulos(eixos, rotulos)

    figura.tight_layout()
    figura.savefig(destino, bbox_inches="tight")
    plt.close(figura)


MAX_ROTULOS_LEGIVEIS = 15
# acima disto as marcas fundem-se numa faixa grossa e escondem a linha
MAX_MARCAS_LEGIVEIS = 40


def rotulo_do_eixo_y(grafico: dict) -> str:
    """O que se escreve no eixo y.

    Com series a partir de varias colunas nao ha uma coluna de valores: dizer
    «Registos» seria errado, porque nao sao contagens de registos.
    """
    if isinstance(grafico.get("serie"), list):
        return ""
    return str(grafico.get("eixo_y") or "Registos")

# cores e tracos distinguiveis tambem a preto e branco e por quem nao distingue
# vermelho de verde: a forma do traco identifica a serie, nao so a cor
CORES = ("#3b6ea5", "#c1553b", "#4a8c5a", "#8a6bbf", "#c99a2e", "#57868c")
TRACOS = ("-", "--", "-.", ":", (0, (3, 1, 1, 1)), (0, (5, 1)))
MARCAS = ("o", "s", "^", "D", "v", "P")


def desenhar_series(eixos, series: dict, categorias, rotulos: list[str],
                    tipo: str) -> None:
    """Desenha varias series no mesmo grafico, com legenda.

    «categorias» e a grelha do eixo x, a do conjunto. Todas as series sao
    alinhadas a ELA e nao a uma grelha propria: se cada uma usasse a sua
    ordem, uma serie que so comece a meio ficaria desenhada por cima dos
    rotulos errados, em silencio.
    """
    nomes = list(series)
    if tipo == "area":
        # empilhadas: ve-se o total e quanto cada serie contribui
        eixos.stackplot(range(len(rotulos)),
                        *[alinhar(series[n], categorias) for n in nomes],
                        labels=nomes, colors=CORES[:len(nomes)], alpha=0.8)
        eixos.set_xticks(range(len(rotulos)))
        eixos.set_xticklabels(rotulos)
        eixos.grid(linestyle=":", alpha=0.6)
        rodar_rotulos(eixos, rotulos)
    elif tipo == "barras_horizontais":
        altura = 0.8 / len(nomes)
        posicoes = range(len(rotulos))
        for posicao, nome in enumerate(nomes):
            deslocamento = (posicao - (len(nomes) - 1) / 2) * altura
            eixos.barh([p + deslocamento for p in posicoes],
                       alinhar(series[nome], categorias),
                       height=altura, label=nome, color=CORES[posicao % len(CORES)])
        eixos.set_yticks(list(posicoes))
        eixos.set_yticklabels(rotulos)
        eixos.invert_yaxis()
        eixos.grid(axis="x", linestyle=":", alpha=0.6)
    elif tipo == "linhas":
        for posicao, nome in enumerate(nomes):
            eixos.plot(rotulos, alinhar(series[nome], categorias), label=nome,
                       linewidth=2, color=CORES[posicao % len(CORES)],
                       linestyle=TRACOS[posicao % len(TRACOS)],
                       marker=(MARCAS[posicao % len(MARCAS)]
                               if len(rotulos) <= MAX_MARCAS_LEGIVEIS else None),
                       markersize=4)
        eixos.grid(linestyle=":", alpha=0.6)
    else:  # barras agrupadas
        largura = 0.8 / len(nomes)
        posicoes = range(len(rotulos))
        for posicao, nome in enumerate(nomes):
            deslocamento = (posicao - (len(nomes) - 1) / 2) * largura
            eixos.bar([p + deslocamento for p in posicoes],
                      alinhar(series[nome], categorias),
                      width=largura, label=nome,
                      color=CORES[posicao % len(CORES)])
        eixos.set_xticks(list(posicoes))
        eixos.set_xticklabels(rotulos)
        eixos.grid(axis="y", linestyle=":", alpha=0.6)

    eixos.legend(fontsize=8, framealpha=0.9)


def alinhar(serie: pd.Series, categorias) -> list[float]:
    """Poe a serie na grelha do eixo x.

    Uma serie pode nao ter todas as categorias -- um canal que so existiu a
    partir de certa altura. Falta e falta: entra como buraco, nao como zero,
    porque zero seria um valor inventado.
    """
    grelha = []
    for categoria in categorias:
        valor = serie.get(categoria)
        grelha.append(float("nan") if valor is None else float(valor))
    return grelha


def rodar_rotulos(eixos, rotulos: list[str]) -> None:
    """Roda os rotulos do eixo x, e desbasta-os quando sao demasiados.

    Com 36 periodos os rotulos sobrepoem-se e o eixo fica ilegivel. Mostra-se
    um em cada N; os pontos ficam todos no grafico, so os rotulos e que
    espacam.
    """
    mais_comprido = max((len(r) for r in rotulos), default=0)
    if len(rotulos) > 8 or mais_comprido > 6:
        for rotulo in eixos.get_xticklabels():
            rotulo.set_rotation(45)
            rotulo.set_horizontalalignment("right")

    if len(rotulos) > MAX_ROTULOS_LEGIVEIS:
        passo = -(-len(rotulos) // MAX_ROTULOS_LEGIVEIS)  # divisao a arredondar para cima
        posicoes = list(range(0, len(rotulos), passo))
        eixos.set_xticks(posicoes)
        eixos.set_xticklabels([rotulos[i] for i in posicoes],
                              rotation=45, horizontalalignment="right")


# ------------------------------------------------------------------- texto


def frase_descritiva(serie: pd.Series, grafico: dict, n_linhas: int) -> str:
    """Frase descritiva, com numeros vindos da mesma conta que desenhou o grafico."""
    eixo_x = grafico["eixo_x"]
    eixo_y = grafico.get("eixo_y")

    if grafico["tipo"] in TIPOS_SEM_AGREGACAO:
        return (
            f"Cada ponto é uma linha do ficheiro: «{eixo_x}» na horizontal e "
            f"«{eixo_y}» na vertical, a partir de {formatar_numero(n_linhas)} "
            "linhas. O gráfico mostra como as duas colunas se relacionam."
        )

    agregacao = grafico["agregacao"]
    if agregacao == "soma":
        o_que = f"o total de «{eixo_y}»"
    elif agregacao == "media":
        o_que = f"a média de «{eixo_y}»"
    elif eixo_y is None:
        o_que = "o número de registos"
    else:
        o_que = f"a contagem de «{eixo_y}»"

    maior_categoria = formatar_categoria(serie.idxmax())
    menor_categoria = formatar_categoria(serie.idxmin())

    partes = [
        f"O gráfico mostra {o_que} por «{eixo_x}», a partir de "
        f"{formatar_numero(n_linhas)} linhas agrupadas em "
        f"{formatar_numero(len(serie))} categorias."
    ]

    if agregacao == "soma":
        partes.append(f"O total é {formatar_numero(serie.sum())}.")
    elif agregacao == "media":
        partes.append(f"A média das categorias é {formatar_numero(serie.mean())}.")
    else:
        partes.append(f"O total de registos é {formatar_numero(serie.sum())}.")

    partes.append(
        f"O valor mais alto é {formatar_numero(serie.max())}, em «{maior_categoria}», "
        f"e o mais baixo é {formatar_numero(serie.min())}, em «{menor_categoria}»."
    )

    if grafico["tipo"] == "circular" and serie.sum() > 0:
        fatia = round(float(serie.max()) / float(serie.sum()) * 100, 1)
        partes.append(
            f"A maior fatia, «{maior_categoria}», representa "
            f"{formatar_numero(fatia)}% do total."
        )

    return " ".join(partes)


# -------------------------------------------------------------- estatistica
# Tudo daqui para baixo trabalha so com listas de numeros, sem dependencias
# novas. Sao funcoes puras de proposito: da para as bater contra uma conta
# feita a parte, que e a unica maneira de verificar estatistica -- um R2
# errado nao se ve a olho como se ve um grafico partido.


def regressao_linear(valores: list[float]) -> dict | None:
    """Minimos quadrados sobre 0, 1, 2...

    Devolve tudo o que a previsao precisa, para nao haver dois sitios a
    calcular a mesma reta: declive, intercecao, r2, erro-padrao dos residuos,
    media dos x e soma dos quadrados dos desvios em x.

    None quando ha poucos pontos para a reta significar alguma coisa.
    """
    n = len(valores)
    if n < MIN_PERIODOS_REGRESSAO:
        return None

    xs = list(range(n))
    media_x = sum(xs) / n
    media_y = sum(valores) / n
    soma_xy = sum((x - media_x) * (y - media_y) for x, y in zip(xs, valores))
    soma_xx = sum((x - media_x) ** 2 for x in xs)
    if soma_xx == 0:
        return None

    declive = soma_xy / soma_xx
    intercecao = media_y - declive * media_x
    residuos = sum((y - (declive * x + intercecao)) ** 2 for x, y in zip(xs, valores))
    variacao_total = sum((y - media_y) ** 2 for y in valores)

    return {
        "declive": declive,
        "intercecao": intercecao,
        "r2": 1.0 if variacao_total == 0 else 1 - residuos / variacao_total,
        "erro_padrao": math.sqrt(residuos / (n - 2)) if n > 2 else 0.0,
        "media_x": media_x,
        "soma_xx": soma_xx,
        "n": n,
    }


def crescimento_medio(valores: list[float]) -> float | None:
    """Taxa geometrica media por periodo, em percentagem.

    Exige valores todos positivos: com um zero pelo meio a raiz nao existe,
    e com negativos o resultado nao quer dizer nada.
    """
    if len(valores) < MIN_PERIODOS_CRESCIMENTO:
        return None
    if any(v <= 0 for v in valores):
        return None
    return ((valores[-1] / valores[0]) ** (1 / (len(valores) - 1)) - 1) * 100


def detetar_atipicos(serie: pd.Series) -> list | None:
    """Valores fora de 1,5 vezes o intervalo interquartil (metodo de Tukey)."""
    if len(serie) < MIN_CATEGORIAS_ATIPICOS:
        return None
    q1 = float(serie.quantile(0.25))
    q3 = float(serie.quantile(0.75))
    intervalo = q3 - q1
    if intervalo == 0:
        return []
    minimo = q1 - 1.5 * intervalo
    maximo = q3 + 1.5 * intervalo
    return [(c, float(v)) for c, v in serie.items() if v < minimo or v > maximo]


def numero_do_mes(rotulo) -> int | None:
    """Tira o numero do mes de «2025-03», de «Marco» ou de uma data."""
    if isinstance(rotulo, pd.Timestamp):
        return int(rotulo.month)
    texto = normalizar(rotulo)
    if texto in MESES_PT:
        # "marco" e "março" sao o mesmo mes; a lista tem os dois
        ordem = ["janeiro", "fevereiro", "marco", "abril", "maio", "junho",
                 "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        return ordem.index(texto.replace("ç", "c")) + 1
    encontrado = re.fullmatch(r"(\d{4})[-/](\d{1,2})", texto)
    if encontrado:
        mes = int(encontrado.group(2))
        return mes if 1 <= mes <= 12 else None
    return None


def indice_sazonal(serie: pd.Series) -> dict | None:
    """Indice sazonal classico: media de cada mes a dividir pela media geral.

    Um indice de 1,4 quer dizer que aquele mes vale, em media, 40% acima do
    mes tipico. Devolve None quando nao ha ciclos completos que cheguem --
    dois anos de dados mensais e o minimo para a media de cada mes significar
    alguma coisa.
    """
    meses = [numero_do_mes(rotulo) for rotulo in serie.index]
    if any(m is None for m in meses):
        return None

    anos = {extrair_ano(rotulo) for rotulo in serie.index}
    anos.discard(None)
    if len(anos) < MIN_CICLOS_SAZONALIDADE or len(serie) < 12 * MIN_CICLOS_SAZONALIDADE:
        return None

    # O indice sazonal e uma RAZAO: so tem significado em dados de escala
    # racional, todos positivos. Encontrado com dados reais de anomalias de
    # temperatura, que oscilam a volta de zero: a media perto de zero fazia o
    # indice explodir e inverter sinais, e saia «indice -0,34 -- 134% abaixo»,
    # que e um disparate com ar de rigor.
    valores = [float(v) for v in serie.values]
    if any(v <= 0 for v in valores):
        return {"impossivel": "negativos"}

    media_geral = float(serie.mean())
    if media_geral <= 0:
        return {"impossivel": "media"}

    por_mes: dict[int, list[float]] = {}
    for mes, valor in zip(meses, serie.values):
        por_mes.setdefault(mes, []).append(float(valor))

    indices = {m: (sum(vs) / len(vs)) / media_geral for m, vs in por_mes.items()}
    return {"indices": indices, "ciclos": len(anos)}


def valor_t(graus: int) -> float:
    """t de Student a 95%, bilateral."""
    if graus <= 0:
        return T_95[1]
    return T_95.get(graus, 1.96)


def proximos_rotulos(indice, quantos: int) -> list[str]:
    """Continua a sequencia de rotulos para os periodos a prever.

    Quando nao consegue continuar a sequencia devolve «período +1», «+2»...
    Inventar um rotulo errado seria pior do que nao ter rotulo nenhum.
    """
    ultimo = indice[-1]

    if isinstance(ultimo, pd.Timestamp) and len(indice) >= 2:
        passo = ultimo - indice[-2]
        return [(ultimo + passo * (i + 1)).strftime("%d/%m/%Y") for i in range(quantos)]

    texto = str(ultimo)
    encontrado = re.fullmatch(r"(\d{4})([-/])(\d{1,2})", texto)
    if encontrado:
        ano, separador, mes = int(encontrado.group(1)), encontrado.group(2), int(encontrado.group(3))
        rotulos = []
        for _ in range(quantos):
            mes += 1
            if mes > 12:
                mes, ano = 1, ano + 1
            rotulos.append(f"{ano}{separador}{mes:02d}")
        return rotulos

    if e_ano(ultimo):
        primeiro = int(como_numero(ultimo))
        return [str(primeiro + i + 1) for i in range(quantos)]

    numero = numero_do_mes(ultimo)
    if numero is not None and normalizar(ultimo) in MESES_PT:
        nomes = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho",
                 "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        return [nomes[(numero - 1 + i + 1) % 12] for i in range(quantos)]

    return [f"período +{i + 1}" for i in range(quantos)]


def prever(serie: pd.Series, horizonte: int) -> dict | None:
    """Previsao por decomposicao: dessazonaliza, ajusta a reta, extrapola.

    A reta e ajustada sobre a serie JA sem sazonalidade -- ajustar sobre os
    dados em bruto daria um R2 pior e uma tendencia contaminada pelo ciclo.
    No fim, o fator sazonal de cada periodo previsto e reaplicado.

    Devolve None quando as guardas nao deixam; quem chama e que escreve a
    razao no relatorio.
    """
    valores = [float(v) for v in serie.values]
    if len(valores) < MIN_PERIODOS_PREVISAO:
        return None

    sazonal = indice_sazonal(serie)
    if sazonal is not None and "impossivel" in sazonal:
        sazonal = None  # sem escala racional nao ha fator sazonal a aplicar
    if sazonal is not None:
        fatores = [sazonal["indices"].get(numero_do_mes(r), 1.0) for r in serie.index]
        if any(f <= 0 for f in fatores):
            sazonal = None

    if sazonal is None:
        base = valores
        metodo = "tendência linear sobre a série observada"
    else:
        base = [v / f for v, f in zip(valores, fatores)]
        metodo = ("tendência linear sobre a série dessazonalizada, com o fator "
                  "sazonal reaplicado a cada período previsto")

    reta = regressao_linear(base)
    if reta is None or reta["r2"] < R2_MINIMO_PREVISAO:
        return {"recusa": "ajuste", "r2": None if reta is None else reta["r2"]}

    n = reta["n"]
    graus = n - 2
    t = valor_t(graus)
    rotulos = proximos_rotulos(serie.index, horizonte)

    previsoes = []
    for passo in range(horizonte):
        x = n + passo
        centro = reta["intercecao"] + reta["declive"] * x
        margem = t * reta["erro_padrao"] * math.sqrt(
            1 + 1 / n + (x - reta["media_x"]) ** 2 / reta["soma_xx"]
        )
        inferior, superior = centro - margem, centro + margem

        if sazonal is not None:
            proximo_mes = numero_do_mes(rotulos[passo])
            fator = sazonal["indices"].get(proximo_mes, 1.0)
            centro, inferior, superior = centro * fator, inferior * fator, superior * fator

        previsoes.append({
            "rotulo": rotulos[passo],
            "centro": centro,
            "inferior": inferior,
            "superior": superior,
        })

    return {
        "previsoes": previsoes,
        "metodo": metodo,
        "r2": reta["r2"],
        "n": n,
        "sazonal": sazonal is not None,
    }


def horizonte_permitido(n_periodos: int) -> int:
    """Quantos periodos e defensavel prever a partir de uma serie de n."""
    return max(1, min(MAX_HORIZONTE_PREVISAO,
                      int(n_periodos * FRACAO_MAXIMA_HORIZONTE)))


def classificar(valor: float, baixo: float, alto: float,
                termos: tuple[str, str, str]) -> str:
    if valor < baixo:
        return termos[0]
    if valor <= alto:
        return termos[1]
    return termos[2]


def criterio(baixo: float, alto: float, termos: tuple[str, str, str],
             sufixo: str = "%") -> str:
    return (f"critério: <{formatar_numero(baixo)}{sufixo} {termos[0]}, "
            f"{formatar_numero(baixo)}–{formatar_numero(alto)}{sufixo} {termos[1]}, "
            f">{formatar_numero(alto)}{sufixo} {termos[2]}")


# ------------------------------------------------------------ eixo temporal


def como_numero(valor) -> float | None:
    """Devolve o valor como float, ou None se nao for numero.

    Existe porque o pandas devolve inteiros do numpy, que nao passam num
    isinstance(valor, int) do Python. Confiar no isinstance dava falsos
    negativos silenciosos.
    """
    if isinstance(valor, bool) or isinstance(valor, pd.Timestamp):
        return None
    try:
        numero = float(valor)
    except (TypeError, ValueError):
        return None
    return None if math.isnan(numero) else numero


def e_ano(valor) -> bool:
    numero = como_numero(valor)
    return numero is not None and numero.is_integer() and ANO_MINIMO <= numero <= ANO_MAXIMO


def partes_de_data(texto: str) -> tuple[int, int, int] | None:
    """Parte «31/01/2026» ou «2026-01-31» em (primeiro, segundo, ano).

    Nao decide ainda o que e dia e o que e mes: so separa. A decisao vem
    depois, e so quando a coluna a provar.
    """
    limpo = str(texto).strip()
    iso = re.fullmatch(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", limpo)
    if iso:
        return (int(iso.group(3)), int(iso.group(2)), int(iso.group(1)))  # ja e ano-mes-dia

    outro = re.fullmatch(r"(\d{1,2})[-/.](\d{1,2})[-/.](\d{4})", limpo)
    if outro:
        return (int(outro.group(1)), int(outro.group(2)), int(outro.group(3)))
    return None


def converter_datas_texto(valores, onde: str, nome: str) -> list | None:
    """Converte uma coluna de datas em texto, mas so com a convencao provada.

    «01/02/2026» e 1 de Fevereiro em Portugal e 2 de Janeiro em ingles. Basta
    um dia acima de 12 algures na coluna para a duvida acabar. Quando nada o
    prova, devolve None e a coluna fica categorica: falhar a analise e seguro,
    trocar Janeiro por Fevereiro nao e -- daria um indice sazonal errado com ar
    de rigor, que e a pior combinacao possivel.
    """
    partes = []
    for valor in valores:
        if valor is None or (isinstance(valor, float) and math.isnan(valor)):
            return None
        separado = partes_de_data(valor)
        if separado is None:
            return None
        partes.append(separado)

    if not partes:
        return None

    # ISO ja vem em (dia, mes, ano) resolvido; para as outras, procurar prova
    dia_primeiro = any(p[0] > 12 for p in partes)
    mes_primeiro = any(p[1] > 12 for p in partes)

    if dia_primeiro and mes_primeiro:
        raise ErroDados(
            f"A coluna «{nome}» usada no {onde} mistura formatos de data "
            "incompatíveis: umas parecem dia/mês e outras mês/dia. "
            "Formata a coluna toda como data no Excel."
        )
    if not dia_primeiro and not mes_primeiro:
        todas_iso = all(
            re.fullmatch(r"\d{4}[-/.]\d{1,2}[-/.]\d{1,2}", str(v).strip())
            for v in valores
        )
        if not todas_iso:
            return None

    convertidas = []
    for primeiro, segundo, ano in partes:
        dia, mes = (primeiro, segundo) if not mes_primeiro else (segundo, primeiro)
        try:
            convertidas.append(pd.Timestamp(year=ano, month=mes, day=dia))
        except ValueError:
            return None
    return convertidas


def parece_temporal(indice) -> bool:
    """Diz se os rotulos do eixo x sao uma linha do tempo.

    Na duvida devolve False. E a hipotese segura: perde-se analise, nao se
    ganham numeros errados. Uma regressao sobre «Canal» mudaria de declive
    so por trocar duas colunas no Excel.
    """
    valores = [v for v in indice if v is not None]
    if not valores:
        return False

    if all(isinstance(v, pd.Timestamp) for v in valores):
        return True
    if all(e_ano(v) for v in valores):
        return True

    textos = [normalizar(v) for v in valores]
    if all(t in MESES_PT for t in textos):
        return True
    # «2023-01», o formato mais comum em exportacoes de plataformas
    if all(numero_do_mes(v) is not None and extrair_ano(v) is not None
           for v in valores):
        return True
    if all(re.fullmatch(r"t[1-4]|[1-4]\.?[ºo]?\s*trimestre", t) for t in textos):
        return True
    return False


def extrair_ano(valor):
    """Tira o ano de uma data, de um numero ou de um texto.

    O ano nao sofre da ambiguidade dd/mm contra mm/dd: em «01/02/2026» o ano
    e o mesmo nas duas leituras. Por isso a quebra por ano e segura mesmo
    quando a ordem do dia e do mes nao esta provada.
    """
    if isinstance(valor, pd.Timestamp):
        return int(valor.year)
    if e_ano(valor):
        return int(como_numero(valor))
    if como_numero(valor) is not None:
        return None
    encontrado = re.search(r"\b\d{4}\b", str(valor))
    return int(encontrado.group(0)) if encontrado else None


# ------------------------------------------------------------------ analise


def secao_meta_por_serie(series: dict, grafico: dict) -> tuple[str, str]:
    """Compara cada serie com a meta."""
    alvo = como_numero(grafico["meta"]["valor"])
    agregacao = grafico["agregacao"]
    totais = {
        nome: (float(s.mean()) if agregacao == "media" else float(s.sum()))
        for nome, s in series.items()
    }
    atingiram = [n for n, v in totais.items() if v >= alvo]
    falharam = sorted(((n, v) for n, v in totais.items() if v < alvo),
                      key=lambda par: par[1])

    texto = (
        f"Meta definida: {formatar_com_precisao(alvo)} por série. "
        f"{formatar_numero(len(atingiram))} de {formatar_numero(len(series))} "
        "séries atingiram-na"
    )
    texto += f" ({listar(atingiram)})." if atingiram else "."

    if falharam:
        lista = "; ".join(
            f"«{n}» com {formatar_com_precisao(v)} "
            f"({formatar_com_precisao(alvo - v)} abaixo)" for n, v in falharam
        )
        texto += f" Abaixo da meta: {lista}."
    return ("Meta", texto)


def correlacao(xs: list[float], ys: list[float]) -> float | None:
    """Coeficiente de correlacao de Pearson, entre -1 e 1."""
    n = len(xs)
    if n < 3:
        return None
    media_x, media_y = sum(xs) / n, sum(ys) / n
    covariancia = sum((x - media_x) * (y - media_y) for x, y in zip(xs, ys))
    desvio_x = math.sqrt(sum((x - media_x) ** 2 for x in xs))
    desvio_y = math.sqrt(sum((y - media_y) ** 2 for y in ys))
    if desvio_x == 0 or desvio_y == 0:
        return None
    return covariancia / (desvio_x * desvio_y)


def analise_dispersao(serie: pd.Series, grafico: dict) -> list[tuple[str, str]]:
    """Analise de um grafico de dispersao: intervalos e correlacao.

    Termina sempre a dizer que correlacao nao e causa. A tentacao de ler
    «investir mais causa mais conversoes» e enorme, e o ficheiro nao tem como
    saber isso -- apontar causas continua proibido.
    """
    xs = [como_numero(c) for c in serie.index]
    ys = [float(v) for v in serie.values]
    eixo_x, eixo_y = grafico["eixo_x"], grafico["eixo_y"]

    blocos = [("Âmbito", (
        f"{formatar_numero(len(ys))} pontos, um por linha do ficheiro. "
        f"«{eixo_x}» vai de {formatar_com_precisao(min(xs))} a "
        f"{formatar_com_precisao(max(xs))}; «{eixo_y}» vai de "
        f"{formatar_com_precisao(min(ys))} a {formatar_com_precisao(max(ys))}."
    ))]

    r = correlacao(xs, ys)
    if r is None:
        blocos.append(("Relação", (
            "Não calculada — uma das colunas tem sempre o mesmo valor, e sem "
            "variação não há relação que medir."
        )))
        return blocos

    termos = ("fraca", "moderada", "forte")
    forca = classificar(abs(r), R_CORRELACAO_FRACA, R_CORRELACAO_FORTE, termos)
    sentido = ("positiva: sobem juntas" if r > 0
               else "negativa: quando uma sobe, a outra desce")
    blocos.append(("Relação", (
        f"Coeficiente de correlação de Pearson de {formatar_numero(round(r, 2))} "
        f"— correlação {forca} e {sentido} "
        f"({criterio(R_CORRELACAO_FRACA, R_CORRELACAO_FORTE, termos, sufixo='')}, "
        "em valor absoluto)."
    )))
    blocos.append(("O que isto não diz", (
        "Correlação não é causa. Duas colunas subirem juntas não significa que "
        "uma faça a outra subir: pode haver um terceiro fator, ou ser "
        "coincidência. Este relatório mede a relação; não a explica."
    )))
    return blocos


def montar_analise(serie: pd.Series, grafico: dict, n_linhas: int,
                   temporal: bool, serie_anual: pd.Series | None
                   ) -> list[tuple[str, str]]:
    """Constroi o bloco de analise como pares (etiqueta, texto).

    Regra unica deste bloco: nenhuma palavra qualitativa entra sem uma regra
    calculada, e a regra vai escrita ao lado. Nunca «bom» ou «fraco» -- isso
    exigiria uma meta que o ficheiro Excel nao tem.
    """
    if grafico["tipo"] in TIPOS_SEM_AGREGACAO:
        return analise_dispersao(serie, grafico)

    blocos: list[tuple[str, str]] = []
    valores = [float(v) for v in serie.values]
    total = sum(valores)
    media = total / len(valores)

    blocos.append(("Âmbito", (
        f"{formatar_numero(len(serie))} categorias de «{grafico['eixo_x']}», "
        f"a partir de {formatar_numero(n_linhas)} linhas. "
        f"Total {formatar_numero(total)}; média {formatar_numero(media)}; "
        f"mediana {formatar_numero(float(serie.median()))}."
    )))

    # o ambito «serie» e tratado uma vez so, ao nivel do grafico
    if grafico.get("meta") and grafico["meta"].get("ambito", "total") != "serie":
        blocos.append(secao_meta(serie, grafico))

    blocos.append(("Amplitude", (
        f"Máximo {formatar_numero(serie.max())} em «{formatar_categoria(serie.idxmax())}»; "
        f"mínimo {formatar_numero(serie.min())} em «{formatar_categoria(serie.idxmin())}». "
        f"Amplitude {formatar_numero(float(serie.max()) - float(serie.min()))}"
        + (f"; rácio de {formatar_numero(round(float(serie.max()) / float(serie.min()), 2))}."
           if float(serie.min()) > 0 else " (o mínimo é zero ou negativo, logo sem rácio).")
    )))

    if len(serie) > 1 and media != 0:
        desvio = float(serie.std(ddof=1))
        cv = abs(desvio / media) * 100
        termos = ("baixa", "moderada", "elevada")
        blocos.append(("Dispersão", (
            f"Desvio-padrão {formatar_numero(desvio)}; coeficiente de variação "
            f"{formatar_numero(round(cv, 1))}% — dispersão "
            f"{classificar(cv, CV_DISPERSAO_BAIXA, CV_DISPERSAO_ELEVADA, termos)} "
            f"({criterio(CV_DISPERSAO_BAIXA, CV_DISPERSAO_ELEVADA, termos)})."
        )))

    if total > 0 and all(v >= 0 for v in valores):
        ordenados = sorted(valores, reverse=True)
        peso_maior = ordenados[0] / total * 100
        termos = ("baixa", "moderada", "elevada")
        texto = (
            f"A maior categoria representa {formatar_numero(round(peso_maior, 1))}% do "
            f"total — concentração "
            f"{classificar(peso_maior, PESO_CONCENTRACAO_BAIXA, PESO_CONCENTRACAO_ELEVADA, termos)} "
            f"({criterio(PESO_CONCENTRACAO_BAIXA, PESO_CONCENTRACAO_ELEVADA, termos)})."
        )
        if len(serie) > 3:
            peso_tres = sum(ordenados[:3]) / total * 100
            texto += f" As três maiores valem {formatar_numero(round(peso_tres, 1))}%."
        blocos.append(("Concentração", texto))

    atipicos = detetar_atipicos(serie)
    if atipicos is None:
        blocos.append(("Valores atípicos", (
            f"Não avaliados — o método de Tukey (1,5×IQR) precisa de pelo menos "
            f"{MIN_CATEGORIAS_ATIPICOS} categorias e há {formatar_numero(len(serie))}."
        )))
    elif atipicos:
        lista = "; ".join(
            f"«{formatar_categoria(c)}» ({formatar_numero(v)})" for c, v in atipicos
        )
        quantos = ("1 valor" if len(atipicos) == 1
                   else f"{formatar_numero(len(atipicos))} valores")
        blocos.append(("Valores atípicos", (
            f"{quantos} fora do intervalo interquartil alargado "
            f"(método de Tukey, 1,5×IQR): {lista}."
        )))
    else:
        blocos.append(("Valores atípicos", (
            "Nenhuma categoria fora do intervalo interquartil alargado "
            "(método de Tukey, 1,5×IQR)."
        )))

    pedido = grafico.get("previsao")

    if temporal:
        blocos.extend(analise_temporal(serie, valores))
        if pedido:
            blocos.append(secao_previsao(serie, pedido))
    else:
        blocos.append(("Evolução", (
            f"Não analisada — «{grafico['eixo_x']}» é um eixo categórico, não uma "
            "linha do tempo. Tendências e regressões só têm significado quando a "
            "ordem das categorias é a ordem do tempo."
        )))
        if pedido:
            blocos.append(("Previsão", (
                f"Não calculada — «{grafico['eixo_x']}» é um eixo categórico. Prever o "
                "período seguinte só faz sentido quando existe um período seguinte."
            )))

    if serie_anual is not None and len(serie_anual) >= 2:
        blocos.extend(analise_anual(serie_anual))

    return blocos


def secao_comparacao_series(series: dict, grafico: dict,
                            temporal: bool) -> list[tuple[str, str]]:
    """Compara as series entre si: peso de cada uma e quem cresceu mais."""
    agregacao = grafico["agregacao"]
    resumo = {
        nome: (float(s.mean()) if agregacao == "media" else float(s.sum()))
        for nome, s in series.items()
    }
    total = sum(resumo.values())

    partes = []
    for nome, valor in resumo.items():
        peso = f", {formatar_numero(round(valor / total * 100, 1))}%" if total > 0 else ""
        partes.append(f"«{nome}» {formatar_com_precisao(valor)}{peso}")

    o_que = "média" if agregacao == "media" else "total"
    origem = ("uma por coluna" if isinstance(grafico["serie"], list)
              else f"de «{grafico['serie']}»")
    texto = (
        f"{formatar_numero(len(series))} séries, {origem}. "
        f"Por {o_que}: " + "; ".join(partes) + "."
    )

    if temporal:
        crescimentos = {}
        for nome, s in series.items():
            valores = [float(v) for v in s.values]
            if len(valores) >= 2 and valores[0] != 0:
                crescimentos[nome] = (valores[-1] - valores[0]) / abs(valores[0]) * 100
        if crescimentos:
            maior = max(crescimentos, key=crescimentos.get)
            menor = min(crescimentos, key=crescimentos.get)
            texto += (
                f" Do primeiro ao último período, «{maior}» variou "
                f"{'+' if crescimentos[maior] >= 0 else ''}"
                f"{formatar_numero(round(crescimentos[maior], 1))}%"
            )
            if menor != maior:
                texto += (
                    f" e «{menor}» {'+' if crescimentos[menor] >= 0 else ''}"
                    f"{formatar_numero(round(crescimentos[menor], 1))}%"
                )
            texto += "."

    blocos = [("Comparação entre séries", texto)]

    linhas = []
    for nome, s in series.items():
        parte = (
            f"«{nome}»: máximo {formatar_com_precisao(s.max())} em "
            f"«{formatar_categoria(s.idxmax())}», mínimo "
            f"{formatar_com_precisao(s.min())} em «{formatar_categoria(s.idxmin())}»"
        )
        if temporal and len(s) >= MIN_PERIODOS_REGRESSAO:
            reta = regressao_linear([float(v) for v in s.values])
            if reta and reta["r2"] >= R2_AJUSTE_FRACO:
                parte += (", tendência crescente" if reta["declive"] > 0
                          else ", tendência decrescente")
            elif reta:
                parte += ", tendência indefinida"
        linhas.append(parte + ".")

    blocos.append(("Por série", " ".join(linhas)))
    return blocos


def secao_meta(serie: pd.Series, grafico: dict) -> tuple[str, str]:
    """Compara os dados com a meta que o utilizador deu.

    Esta e a unica seccao que avalia desempenho, e so existe porque a
    referencia vem de fora: sem meta, «fraco» seria uma opiniao; com meta,
    «atingiu 91%» e uma conta. Continua sem dizer se o resultado e bom.
    """
    meta = grafico["meta"]
    alvo = como_numero(meta["valor"])
    ambito = meta.get("ambito", "total")
    agregacao = grafico["agregacao"]

    if ambito == "total":
        if agregacao == "media":
            observado = float(serie.mean())
            o_que = "A média das categorias"
        else:
            observado = float(serie.sum())
            o_que = "O total"

        if alvo == 0:
            return ("Meta", (
                f"{o_que} é {formatar_com_precisao(observado)}, face a uma meta de "
                "zero. Sem meta positiva não há percentagem a calcular."
            ))

        percentagem = observado / alvo * 100
        diferenca = observado - alvo
        posicao = ("acima da meta" if diferenca > 0
                   else "abaixo da meta" if diferenca < 0 else "exatamente na meta")
        texto = (
            f"Meta definida: {formatar_com_precisao(alvo)}. "
            f"{o_que} é {formatar_com_precisao(observado)} — "
            f"{formatar_numero(round(percentagem, 1))}% da meta, "
            f"{formatar_com_precisao(abs(diferenca))} {posicao}."
        )
        return ("Meta", texto)

    # ambito == "categoria"
    valores = [float(v) for v in serie.values]
    atingiram = [c for c, v in serie.items() if float(v) >= alvo]
    falharam = [(c, float(v)) for c, v in serie.items() if float(v) < alvo]

    texto = (
        f"Meta definida: {formatar_com_precisao(alvo)} por categoria. "
        f"{formatar_numero(len(atingiram))} de {formatar_numero(len(serie))} "
        f"categorias atingiram-na "
        f"({formatar_numero(round(len(atingiram) / len(serie) * 100, 1))}%)."
    )

    if falharam:
        piores = sorted(falharam, key=lambda par: par[1])[:3]
        lista = "; ".join(
            f"«{formatar_categoria(c)}» com {formatar_com_precisao(v)} "
            f"({formatar_com_precisao(alvo - v)} abaixo)"
            for c, v in piores
        )
        texto += f" Mais distantes da meta: {lista}."
        em_falta = sum(alvo - v for _, v in falharam)
        texto += (
            f" Somando o que faltou a cada uma, o défice total é "
            f"{formatar_com_precisao(em_falta)}."
        )
    else:
        texto += " Nenhuma categoria ficou abaixo."

    media = sum(valores) / len(valores)
    texto += (
        f" A média por categoria é {formatar_com_precisao(media)}, "
        f"{formatar_numero(round(media / alvo * 100, 1))}% da meta."
        if alvo != 0 else ""
    )
    return ("Meta", texto)


def secao_previsao(serie: pd.Series, pedido: int) -> tuple[str, str]:
    """Constroi a seccao Previsao, ou a explicacao de porque nao a ha."""
    n = len(serie)
    if n < MIN_PERIODOS_PREVISAO:
        return ("Previsão", (
            f"Não calculada — são precisos pelo menos {MIN_PERIODOS_PREVISAO} "
            f"períodos observados e há {formatar_numero(n)}."
        ))

    maximo = horizonte_permitido(n)
    horizonte = min(pedido, maximo)
    aviso_horizonte = ""
    if pedido > maximo:
        aviso_horizonte = (
            f" Pediste {formatar_numero(pedido)} períodos, mas o máximo defensável "
            f"para uma série de {formatar_numero(n)} é {formatar_numero(maximo)} "
            f"(um terço da série, no máximo {MAX_HORIZONTE_PREVISAO}); "
            "extrapolar mais longe seria ficção."
        )

    resultado = prever(serie, horizonte)
    if resultado is None:
        return ("Previsão", f"Não calculada — série demasiado curta.{aviso_horizonte}")

    if "recusa" in resultado:
        r2 = resultado["r2"]
        detalhe = (f"o ajuste da tendência é fraco (R² de {formatar_numero(round(r2, 2))}, "
                   f"mínimo {formatar_numero(R2_MINIMO_PREVISAO)})"
                   if r2 is not None else "não foi possível ajustar uma tendência")
        return ("Previsão", (
            f"Não calculada — {detalhe}. Extrapolar uma reta que não descreve os "
            "dados dava um número com ar de rigor e sem rigor nenhum."
        ))

    linhas = "; ".join(
        f"{p['rotulo']} entre {formatar_com_precisao(p['inferior'])} e "
        f"{formatar_com_precisao(p['superior'])}"
        for p in resultado["previsoes"]
    )

    texto = (
        f"Método: {resultado['metodo']}, ajustada a "
        f"{formatar_numero(resultado['n'])} períodos (R² de "
        f"{formatar_numero(round(resultado['r2'], 2))}). "
        f"Intervalos a {CONFIANCA_PREVISAO}% — {linhas}."
        f"{aviso_horizonte}"
    )

    negativas = [p for p in resultado["previsoes"] if p["centro"] < 0]
    if negativas and all(v >= 0 for v in serie.values):
        texto += (
            f" Atenção: {formatar_numero(len(negativas))} período(s) previstos descem "
            "abaixo de zero, apesar de nenhum valor observado ser negativo — sinal de "
            "que a reta deixou de ser adequada neste horizonte."
        )

    largos = [p for p in resultado["previsoes"]
              if abs(p["centro"]) > 0 and (p["superior"] - p["inferior"]) > abs(p["centro"])]
    if largos:
        texto += (
            f" Em {formatar_numero(len(largos))} período(s) o intervalo é mais largo do "
            "que o próprio valor previsto: a incerteza é da ordem de grandeza da previsão."
        )

    texto += (
        " Pressupõe que a tendência e a sazonalidade observadas se mantêm. Não incorpora "
        "nada que não esteja no ficheiro."
    )
    return ("Previsão", texto)


def analise_temporal(serie: pd.Series, valores: list[float]) -> list[tuple[str, str]]:
    """Secoes que so fazem sentido quando o eixo e uma linha do tempo."""
    blocos = []
    primeiro, ultimo = valores[0], valores[-1]
    variacao = ultimo - primeiro
    texto = (
        f"Do primeiro ao último período: {'+' if variacao >= 0 else ''}"
        f"{formatar_numero(variacao)}"
    )
    if primeiro != 0:
        texto += f" ({'+' if variacao >= 0 else ''}{formatar_numero(round(variacao / abs(primeiro) * 100, 1))}%)"
    texto += "."

    diferencas = [b - a for a, b in zip(valores, valores[1:])]
    subidas = sum(1 for d in diferencas if d > 0)
    descidas = sum(1 for d in diferencas if d < 0)
    texto += (
        f" Das {formatar_numero(len(diferencas))} variações período a período, "
        f"{formatar_numero(subidas)} positivas e {formatar_numero(descidas)} negativas"
    )
    # os zeros contam: uma serie com subidas e patamares, sem descidas, e
    # monotona nao decrescente. Chamar-lhe «nao monotona» era errado.
    if not diferencas:
        pass
    elif descidas == 0 and subidas:
        texto += (" — série monótona crescente." if subidas == len(diferencas)
                  else " — série monótona não decrescente (há períodos iguais).")
    elif subidas == 0 and descidas:
        texto += (" — série monótona decrescente." if descidas == len(diferencas)
                  else " — série monótona não crescente (há períodos iguais).")
    else:
        texto += " — série não monótona."

    if diferencas:
        maior = max(range(len(diferencas)), key=lambda i: diferencas[i])
        menor = min(range(len(diferencas)), key=lambda i: diferencas[i])
        rotulos = [formatar_categoria(c) for c in serie.index]
        texto += (
            f" Maior subida: «{rotulos[maior]}» → «{rotulos[maior + 1]}», "
            f"+{formatar_numero(diferencas[maior])}."
            f" Maior descida: «{rotulos[menor]}» → «{rotulos[menor + 1]}», "
            f"{formatar_numero(diferencas[menor])}."
        )
    blocos.append(("Evolução", texto))

    resultado = regressao_linear(valores)
    if resultado is None:
        blocos.append(("Tendência", (
            f"Regressão não calculada — são precisos pelo menos "
            f"{MIN_PERIODOS_REGRESSAO} períodos e há {formatar_numero(len(valores))}."
        )))
    else:
        declive, r2 = resultado["declive"], resultado["r2"]
        termos = ("fraco", "moderado", "forte")
        ajuste = classificar(r2, R2_AJUSTE_FRACO, R2_AJUSTE_FORTE, termos)
        if r2 < R2_AJUSTE_FRACO:
            leitura = "tendência indefinida: os dados não se aproximam de uma reta"
        elif declive > 0:
            leitura = "tendência crescente"
        elif declive < 0:
            leitura = "tendência decrescente"
        else:
            leitura = "tendência estável"
        blocos.append(("Tendência", (
            f"Regressão linear com declive de {'+' if declive >= 0 else ''}"
            f"{formatar_com_precisao(declive)} por período e R² de "
            f"{formatar_numero(round(r2, 3))} — ajuste {ajuste} "
            f"({criterio(R2_AJUSTE_FRACO, R2_AJUSTE_FORTE, termos, sufixo='')}). "
            f"Leitura: {leitura}."
        )))

    taxa = crescimento_medio(valores)
    if taxa is None:
        if len(valores) < MIN_PERIODOS_CRESCIMENTO:
            razao = (f"são precisos pelo menos {MIN_PERIODOS_CRESCIMENTO} períodos e "
                     f"há {formatar_numero(len(valores))}")
        else:
            razao = "a série tem valores nulos ou negativos, e a taxa geométrica não existe"
        blocos.append(("Crescimento médio", f"Não calculado — {razao}."))
    else:
        blocos.append(("Crescimento médio", (
            f"Taxa de crescimento média geométrica de {'+' if taxa >= 0 else ''}"
            f"{formatar_numero(round(taxa, 1))}% por período."
        )))

    sazonal = indice_sazonal(serie)
    if sazonal is not None and "impossivel" in sazonal:
        blocos.append(("Sazonalidade", (
            "Não avaliada — o índice sazonal é uma razão entre a média de cada mês "
            "e a média geral, e só tem significado quando todos os valores são "
            "positivos. Esta série tem valores nulos ou negativos, e o índice daria "
            "um número sem sentido."
        )))
    elif sazonal is None:
        blocos.append(("Sazonalidade", (
            f"Não avaliada — o índice sazonal precisa de pelo menos "
            f"{MIN_CICLOS_SAZONALIDADE} ciclos anuais completos de dados mensais, "
            "e não foi possível identificá-los nesta série."
        )))
    else:
        nomes = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho",
                 "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        indices = sazonal["indices"]
        mais_forte = max(indices, key=indices.get)
        mais_fraco = min(indices, key=indices.get)
        blocos.append(("Sazonalidade", (
            f"Índice sazonal sobre {formatar_numero(sazonal['ciclos'])} ciclos anuais "
            f"(média de cada mês a dividir pela média geral). "
            f"Mês mais forte: {nomes[mais_forte - 1]}, índice "
            f"{formatar_numero(round(indices[mais_forte], 2))} — "
            f"{formatar_numero(round((indices[mais_forte] - 1) * 100, 1))}% acima do mês "
            f"típico. Mês mais fraco: {nomes[mais_fraco - 1]}, índice "
            f"{formatar_numero(round(indices[mais_fraco], 2))} — "
            f"{formatar_numero(round((1 - indices[mais_fraco]) * 100, 1))}% abaixo."
        )))
    return blocos


def analise_anual(serie_anual: pd.Series) -> list[tuple[str, str]]:
    """Quebra ano a ano, com variacao interanual."""
    anos = list(serie_anual.index)
    valores = [float(v) for v in serie_anual.values]

    partes = [
        f"{ano}: {formatar_numero(valor)}" for ano, valor in zip(anos, valores)
    ]
    blocos = [("Ano a ano", "Totais por ano — " + "; ".join(partes) + ".")]

    variacoes = []
    for anterior, seguinte, valor_a, valor_s in zip(anos, anos[1:], valores, valores[1:]):
        diferenca = valor_s - valor_a
        texto = f"{anterior}→{seguinte}: {'+' if diferenca >= 0 else ''}{formatar_numero(diferenca)}"
        if valor_a != 0:
            texto += f" ({'+' if diferenca >= 0 else ''}{formatar_numero(round(diferenca / abs(valor_a) * 100, 1))}%)"
        variacoes.append(texto)

    positivas = sum(1 for a, b in zip(valores, valores[1:]) if b > a)
    acumulada = valores[-1] - valores[0]
    resumo = (
        "Variação interanual — " + "; ".join(variacoes) + ". "
        f"{formatar_numero(positivas)} de {formatar_numero(len(valores) - 1)} "
        "variações foram positivas. "
        f"Do primeiro ao último ano: {'+' if acumulada >= 0 else ''}"
        f"{formatar_numero(acumulada)}"
    )
    if valores[0] != 0:
        resumo += f" ({'+' if acumulada >= 0 else ''}{formatar_numero(round(acumulada / abs(valores[0]) * 100, 1))}%)"
    resumo += (
        f". Ano mais alto: {serie_anual.idxmax()} "
        f"({formatar_numero(serie_anual.max())}); mais baixo: {serie_anual.idxmin()} "
        f"({formatar_numero(serie_anual.min())})."
    )
    blocos.append(("Comparação entre anos", resumo))
    return blocos


# ---------------------------------------------------------------- documento


def montar_documento(plano: dict, resultados: list[dict], avisos: list[str],
                     notas: list[str], saida: Path,
                     origem: Path | None = None) -> None:
    documento = Document()
    numerar_paginas(documento)
    pagina_de_rosto(documento, plano, origem)
    documento.add_page_break()
    indice(documento, resultados)

    for resultado in resultados:
        documento.add_page_break()

        grafico = resultado["grafico"]
        series = resultado["series"]
        conjunto = resultado["conjunto"]
        varias = [n for n in series if n != ""]

        documento.add_heading(grafico["titulo"], level=1)
        documento.add_picture(str(resultado["imagem"]), width=LARGURA_IMAGEM)
        documento.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        documento.add_paragraph(
            frase_descritiva(conjunto, grafico, resultado["n_linhas"])
        )

        nota = grafico.get("nota")
        if nota:
            paragrafo = documento.add_paragraph()
            paragrafo.add_run(nota).italic = True

        profundidade = plano.get("analise", "completa")
        if profundidade in ("completa", "detalhada"):
            documento.add_heading("Análise", level=2)
            blocos = []
            if varias:
                blocos.extend(secao_comparacao_series(
                    series, grafico, resultado["temporal"]))
                if (grafico.get("meta")
                        and grafico["meta"].get("ambito") == "serie"):
                    blocos.append(secao_meta_por_serie(series, grafico))
            blocos.extend(montar_analise(
                conjunto, grafico, resultado["n_linhas"],
                resultado["temporal"], resultado["serie_anual"],
            ))
            escrever_blocos(documento, blocos)

            if varias and profundidade == "detalhada":
                for nome, serie in series.items():
                    documento.add_heading(f"Análise — {nome}", level=2)
                    escrever_blocos(documento, montar_analise(
                        serie, grafico, len(serie), resultado["temporal"], None))

        if grafico.get("tabela_dados"):
            acrescentar_tabela(documento, conjunto, grafico, notas)

    documento.add_page_break()
    documento.add_heading("Notas sobre os dados", level=1)
    if avisos:
        documento.add_paragraph(
            f"Foram encontrados {len(avisos)} problema(s) nos dados de origem:"
        )
        for linha in avisos:
            documento.add_paragraph(linha, style="List Bullet")
    else:
        documento.add_paragraph("Não foram encontrados problemas nos dados de origem.")

    if notas:
        documento.add_paragraph("Como os dados foram agrupados:")
        for linha in notas:
            documento.add_paragraph(linha, style="List Bullet")

    saida.parent.mkdir(parents=True, exist_ok=True)
    try:
        documento.save(saida)
    except PermissionError:
        raise ErroDados(
            f"Não consigo escrever «{saida}»: o ficheiro está aberto noutro programa "
            "(provavelmente o Word). Fecha-o e tenta outra vez."
        ) from None


def pagina_de_rosto(documento, plano: dict, origem: Path | None) -> None:
    """Titulo, data e ficheiro de origem, numa pagina so."""
    for _ in range(4):  # empurra o titulo para baixo do topo da folha
        documento.add_paragraph()

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = titulo.add_run(plano["titulo_relatorio"])
    corrida.bold = True
    corrida.font.size = Pt(28)

    data = documento.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hoje = datetime.date.today().strftime("%d/%m/%Y")
    corrida = data.add_run(f"Relatório gerado em {hoje}")
    corrida.font.size = Pt(11)

    if origem is not None:
        fonte = documento.add_paragraph()
        fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
        corrida = fonte.add_run(f"A partir de {origem.name}")
        corrida.italic = True
        corrida.font.size = Pt(10)


def indice(documento, resultados: list[dict]) -> None:
    """Lista dos graficos, pela ordem em que aparecem.

    Sem numeros de pagina de proposito: um indice a serio do Word e um campo
    que aparece vazio ate alguem carregar em «atualizar». Uma lista simples
    esta sempre certa.
    """
    documento.add_heading("Índice", level=1)
    for posicao, resultado in enumerate(resultados, start=1):
        # numero OU marca, nunca os dois: «• 1. Titulo» le-se mal
        documento.add_paragraph(f"{posicao}.  {resultado['grafico']['titulo']}")


def numerar_paginas(documento) -> None:
    """Numero de pagina no rodape, centrado.

    O python-docx nao tem isto pronto: mete-se o campo PAGE a mao no XML.
    A capa fica sem numero (primeira pagina diferente).
    """
    seccao = documento.sections[0]
    seccao.different_first_page_header_footer = True

    paragrafo = seccao.footer.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = paragrafo.add_run()
    corrida.font.size = Pt(9)

    abrir = OxmlElement("w:fldChar")
    abrir.set(qn("w:fldCharType"), "begin")
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = "PAGE"
    fechar = OxmlElement("w:fldChar")
    fechar.set(qn("w:fldCharType"), "end")
    for elemento in (abrir, instrucao, fechar):
        corrida._r.append(elemento)


def escrever_blocos(documento, blocos: list[tuple[str, str]]) -> None:
    for etiqueta, texto in blocos:
        paragrafo = documento.add_paragraph()
        paragrafo.add_run(f"{etiqueta}. ").bold = True
        paragrafo.add_run(texto)


def acrescentar_tabela(documento, serie: pd.Series, grafico: dict,
                       notas: list[str]) -> None:
    mostradas = serie.head(MAX_LINHAS_TABELA)
    escondidas = len(serie) - len(mostradas)

    tabela = documento.add_table(rows=1, cols=2)
    tabela.style = "Light Grid Accent 1"
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = str(grafico["eixo_x"])
    cabecalho[1].text = str(grafico.get("eixo_y") or "Registos")

    for categoria, valor in mostradas.items():
        celulas = tabela.add_row().cells
        celulas[0].text = formatar_categoria(categoria)
        celulas[1].text = formatar_numero(valor)

    if escondidas > 0:
        aviso = (
            f"{grafico['titulo']}: a tabela mostra as primeiras "
            f"{MAX_LINHAS_TABELA} de {len(serie)} categorias; "
            f"{escondidas} ficaram de fora."
        )
        notas.append(aviso)
        paragrafo = documento.add_paragraph()
        corrida = paragrafo.add_run(aviso)
        corrida.italic = True
        corrida.font.size = Pt(9)


# --------------------------------------------------------------------- main


def executar(args) -> int:
    plano = ler_plano(Path(args.plano))
    fonte = abrir_dados(Path(args.dados))

    avisos: list[str] = []
    notas: list[str] = list(fonte.notas)
    preparados = []

    for indice, grafico in enumerate(plano["graficos"], start=1):
        series, n_linhas, temporal, anual, conjunto = preparar_dados(
            fonte, grafico, indice, avisos, notas
        )
        preparados.append({
            "grafico": grafico, "series": series, "conjunto": conjunto,
            "n_linhas": n_linhas, "temporal": temporal, "serie_anual": anual,
        })

    print(f"Plano lido: {len(preparados)} gráfico(s).")
    for nota in notas:
        print(f"  · {nota}")

    if avisos:
        print(f"\n{len(avisos)} aviso(s) sobre os dados:")
        for aviso in avisos:
            print(f"  ! {aviso}")
    else:
        print("\nSem avisos: os dados estão limpos.")

    if args.verificar:
        print("\n--verificar: nada foi gerado.")
        return 0

    if avisos and not args.avisos_aceites:
        print(
            "\nO relatório NÃO foi gerado por causa dos avisos acima.\n"
            "Se estiverem todos entendidos e quiseres avançar mesmo assim, "
            "repete o comando com --avisos-aceites."
        )
        return 2

    saida = Path(args.saida)
    pasta_temporaria = Path(tempfile.mkdtemp(prefix="excel-para-word-"))
    try:
        for posicao, resultado in enumerate(preparados, start=1):
            imagem = pasta_temporaria / f"grafico_{posicao}.png"
            desenhar(resultado["conjunto"], resultado["series"],
                     resultado["grafico"], imagem)
            resultado["imagem"] = imagem
        montar_documento(plano, preparados, avisos, notas, saida,
                         origem=Path(args.dados))
    finally:
        shutil.rmtree(pasta_temporaria, ignore_errors=True)

    print(f"\nRelatório gerado: {saida}")
    return 0


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        # a consola do Windows nao usa UTF-8 por omissao e engasga-se nos acentos
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    analisador = argparse.ArgumentParser(
        description="Gera um relatório Word com gráficos a partir de um Excel."
    )
    analisador.add_argument("--dados", required=True, help="ficheiro .xlsx de origem")
    analisador.add_argument("--plano", required=True, help="ficheiro plano.json")
    analisador.add_argument("--saida", help="ficheiro .docx a criar")
    analisador.add_argument(
        "--verificar", action="store_true",
        help="analisa os dados e mostra os problemas, sem gerar nada",
    )
    analisador.add_argument(
        "--avisos-aceites", dest="avisos_aceites", action="store_true",
        help="gera o relatório mesmo havendo avisos",
    )
    args = analisador.parse_args()

    if not args.verificar and not args.saida:
        analisador.error("é preciso --saida (ou usa --verificar para só analisar).")

    try:
        return executar(args)
    except ErroDados as erro:
        print(f"\nErro: {erro}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
