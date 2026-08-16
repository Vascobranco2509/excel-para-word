"""Testes com ficheiros propositadamente sujos.

O que se testa aqui nao e o caso bonito: e o Excel real, com texto onde
deviam estar numeros, celulas vazias, folhas que nao existem e meses
fora de ordem alfabetica.
"""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

import pandas as pd
import pytest

RAIZ = Path(__file__).resolve().parent.parent
SCRIPT = RAIZ / "scripts" / "gerar_relatorio.py"
sys.path.insert(0, str(SCRIPT.parent))

import gerar_relatorio as gr  # noqa: E402

MESES = ["Janeiro", "Fevereiro", "Marco", "Abril", "Maio", "Junho",
         "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]


# ------------------------------------------------------------------ auxiliares


def escrever_excel(caminho: Path, dados: dict, folha: str = "Vendas") -> Path:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(dados).to_excel(caminho, sheet_name=folha, index=False)
    return caminho


def escrever_plano(caminho: Path, graficos: list[dict],
                   titulo: str = "Relatório de teste") -> Path:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    caminho.write_text(
        json.dumps({"titulo_relatorio": titulo, "graficos": graficos}, ensure_ascii=False),
        encoding="utf-8",
    )
    return caminho


def grafico(**alteracoes) -> dict:
    base = {
        "tipo": "barras",
        "folha": "Vendas",
        "eixo_x": "Mes",
        "eixo_y": "Valor",
        "agregacao": "soma",
        "titulo": "Teste",
    }
    base.update(alteracoes)
    return base


def preparar(caminho: Path, config: dict):
    """Corre a preparacao de dados e devolve (serie, n_linhas, avisos, notas)."""
    avisos: list[str] = []
    notas: list[str] = []
    excel = gr.abrir_dados(caminho)
    _, n_linhas, _, _, conjunto = gr.preparar_dados(excel, config, 1, avisos, notas)
    return conjunto, n_linhas, avisos, notas


def preparar_series(caminho: Path, config: dict):
    """Devolve o dicionario {nome: serie} e os avisos."""
    avisos: list[str] = []
    notas: list[str] = []
    fonte = gr.abrir_dados(caminho)
    series, _, temporal, _, _ = gr.preparar_dados(fonte, config, 1, avisos, notas)
    return series, temporal, avisos, notas


def preparar_completo(caminho: Path, config: dict):
    """Como preparar(), mas devolve tambem o eixo temporal e a serie anual."""
    avisos: list[str] = []
    notas: list[str] = []
    excel = gr.abrir_dados(caminho)
    _, n_linhas, temporal, anual, conjunto = gr.preparar_dados(excel, config, 1, avisos, notas)
    return conjunto, n_linhas, temporal, anual, avisos


# --------------------------------------------------------- ordem das categorias


def test_ordem_dos_meses_segue_o_ficheiro_e_nao_o_alfabeto(tmp_path):
    """O erro classico: agrupar por ordem alfabetica poe Abril antes de Janeiro."""
    ficheiro = escrever_excel(
        tmp_path / "meses.xlsx",
        {"Mes": MESES, "Valor": list(range(1, 13))},
    )
    serie, _, avisos, _ = preparar(ficheiro, grafico())

    assert list(serie.index) == MESES
    assert list(serie.index) != sorted(MESES)
    assert avisos == []


# ------------------------------------------------------------- colunas de texto


def test_coluna_de_texto_convertivel_avisa_e_nao_soma_em_silencio(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "texto.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"], "Valor": ["10", "20", "trinta"]},
    )
    serie, n_linhas, avisos, _ = preparar(ficheiro, grafico())

    assert serie.sum() == 30  # "trinta" nao entrou na conta
    assert n_linhas == 2
    assert any("texto" in aviso for aviso in avisos)
    assert any("trinta" in aviso for aviso in avisos)


def test_coluna_sem_numeros_nenhuns_e_erro(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "so_texto.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": ["muito", "pouco"]},
    )
    with pytest.raises(gr.ErroDados, match="não tem números"):
        preparar(ficheiro, grafico())


def test_coluna_de_datas_como_valor_e_erro(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "datas.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"],
         "Valor": pd.to_datetime(["2025-01-01", "2025-02-01"])},
    )
    with pytest.raises(gr.ErroDados, match="datas"):
        preparar(ficheiro, grafico())


# ------------------------------------------------------- numeros com formato
# Encontrados no teste com um ficheiro a serio: uma exportacao portuguesa
# grava o dinheiro como texto, «1.250,00 €», que nao e numero nenhum para
# o pandas.


def test_moeda_a_portuguesa_e_convertida(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "euros.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"],
         "Valor": ["1.250,00 €", "980,50 €", "0,00 €"]},
    )
    serie, _, avisos, _ = preparar(ficheiro, grafico())

    assert serie["Janeiro"] == pytest.approx(1250.0)
    assert serie["Fevereiro"] == pytest.approx(980.5)
    assert serie["Marco"] == pytest.approx(0.0)
    assert any("moeda" in aviso for aviso in avisos)


def test_moeda_a_inglesa_e_convertida(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "dolares.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": ["$1,250.00", "$980.50"]},
    )
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie["Janeiro"] == pytest.approx(1250.0)
    assert serie["Fevereiro"] == pytest.approx(980.5)


def escrever_excel_com_texto(caminho: Path, cabecalhos: list[str],
                             linhas: list[list]) -> Path:
    """Escreve celulas tal e qual, sem o pandas pelo meio a interpretar."""
    from openpyxl import Workbook

    caminho.parent.mkdir(parents=True, exist_ok=True)
    livro = Workbook()
    folha = livro.active
    folha.title = "Vendas"
    folha.append(cabecalhos)
    for linha in linhas:
        folha.append(linha)
    livro.save(caminho)
    return caminho


def test_formato_ambiguo_para_com_erro_em_vez_de_adivinhar(tmp_path):
    """«1.250» tanto pode ser 1250 como 1,25. Adivinhar da um erro invisivel."""
    ficheiro = escrever_excel_com_texto(
        tmp_path / "ambiguo.xlsx",
        ["Mes", "Valor"],
        [["Janeiro", "1.250"], ["Fevereiro", "2.500"]],
    )
    with pytest.raises(gr.ErroDados, match="ambíguo"):
        preparar(ficheiro, grafico())


def test_texto_ambiguo_nao_e_lido_como_decimal_pelo_pandas(tmp_path):
    """O pandas le o texto «1.250» como 1,25. Num Excel portugues isso e mil
    vezes menos. A celula em bruto tem de mandar."""
    ficheiro = escrever_excel_com_texto(
        tmp_path / "mil_vezes.xlsx",
        ["Mes", "Valor"],
        [["Janeiro", "1.250"], ["Fevereiro", "980,50"]],
    )
    serie, _, avisos, _ = preparar(ficheiro, grafico())

    assert serie["Janeiro"] == pytest.approx(1250.0), "leu 1,25 em vez de 1.250"
    assert serie["Fevereiro"] == pytest.approx(980.5)
    assert any("texto" in aviso for aviso in avisos)


def test_numeros_a_serio_nao_sao_mexidos(tmp_path):
    """Uma coluna de numeros verdadeiros nao pode passar pelo caminho do texto."""
    ficheiro = escrever_excel(
        tmp_path / "numeros.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"], "Valor": [1.234, 2.5, 1250.0]},
    )
    serie, _, avisos, _ = preparar(ficheiro, grafico())
    assert serie["Janeiro"] == pytest.approx(1.234)
    assert serie["Marco"] == pytest.approx(1250.0)
    assert avisos == []


def test_valor_ambiguo_resolve_se_a_coluna_o_provar(tmp_path):
    """«980,50» prova que a virgula e decimal; logo «1.250» e mil duzentos e cinquenta."""
    ficheiro = escrever_excel(
        tmp_path / "resolvido.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": ["1.250", "980,50"]},
    )
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie["Janeiro"] == pytest.approx(1250.0)
    assert serie["Fevereiro"] == pytest.approx(980.5)


def test_formatos_misturados_param_com_erro(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "misturado.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": ["1.250,00", "2,500.00"]},
    )
    with pytest.raises(gr.ErroDados, match="incompatíveis"):
        preparar(ficheiro, grafico())


def test_espaco_como_separador_de_milhares(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "espacos.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": ["1 250,00", "980,50"]},
    )
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie["Janeiro"] == pytest.approx(1250.0)


# --------------------------------------------------------------- celulas vazias


def test_celulas_vazias_sao_ignoradas_com_aviso(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "vazias.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", None, "Abril"], "Valor": [10, None, 30, 40]},
    )
    serie, n_linhas, avisos, _ = preparar(ficheiro, grafico())

    assert n_linhas == 2  # so Janeiro e Abril tem os dois campos preenchidos
    assert serie.sum() == 50
    assert any("ignoradas" in aviso for aviso in avisos)


def test_ficheiro_todo_vazio_depois_da_limpeza_e_erro(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "tudo_vazio.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": [None, None]},
    )
    with pytest.raises(gr.ErroDados, match="ficou sem dados"):
        preparar(ficheiro, grafico())


# ------------------------------------------------------------ folhas e colunas


def test_folha_inexistente_lista_as_que_existem(tmp_path):
    ficheiro = escrever_excel(tmp_path / "f.xlsx", {"Mes": ["Janeiro"], "Valor": [1]})
    with pytest.raises(gr.ErroDados) as erro:
        preparar(ficheiro, grafico(folha="Faturacao"))
    assert "Vendas" in str(erro.value)


def test_coluna_inexistente_lista_as_que_existem(tmp_path):
    ficheiro = escrever_excel(tmp_path / "c.xlsx", {"Mes": ["Janeiro"], "Valor": [1]})
    with pytest.raises(gr.ErroDados) as erro:
        preparar(ficheiro, grafico(eixo_y="Receita"))
    assert "Valor" in str(erro.value)


def test_colunas_fantasma_geram_aviso(tmp_path):
    caminho = tmp_path / "fantasma.xlsx"
    df = pd.DataFrame({"Mes": ["Janeiro", "Fevereiro"], "Valor": [10, 20]})
    df["Unnamed: 5"] = [None, "nota solta"]
    df.to_excel(caminho, sheet_name="Vendas", index=False)

    _, _, avisos, _ = preparar(caminho, grafico())
    assert any("sem cabeçalho" in aviso for aviso in avisos)


# -------------------------------------------------------- linhas de totais
# Encontrados num teste com um ficheiro de exportacao a serio: a linha TOTAL
# do fim da folha virava uma categoria e o relatorio contava tudo duas vezes,
# sem avisar.


def test_linha_total_no_fim_da_folha_e_apanhada(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "com_total.xlsx",
        {"Campanha": ["Inverno", "Primavera", "Verao", "TOTAL"],
         "Valor": [100, 200, 300, 600]},
    )
    _, _, avisos, _ = preparar(ficheiro, grafico(eixo_x="Campanha"))
    assert any("totais" in aviso for aviso in avisos)


def test_total_com_acentos_e_maiusculas_tambem_e_apanhado(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "totais.xlsx",
        {"Campanha": ["Inverno", "Primavera", "Total Geral"],
         "Valor": [100, 200, 300]},
    )
    _, _, avisos, _ = preparar(ficheiro, grafico(eixo_x="Campanha"))
    assert any("totais" in aviso for aviso in avisos)


def test_dados_legitimos_nao_sao_confundidos_com_totais(tmp_path):
    """10, 20 e 30 sao dados a serio: o 30 e a soma dos outros por acaso."""
    ficheiro = escrever_excel(
        tmp_path / "legitimo.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"], "Valor": [10, 20, 30]},
    )
    _, _, avisos, _ = preparar(ficheiro, grafico())
    assert avisos == []


def test_aviso_repetido_so_aparece_uma_vez(tmp_path):
    """Dois graficos sobre a mesma folha suja davam o mesmo aviso duas vezes."""
    ficheiro = escrever_excel(
        tmp_path / "suja.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"],
         "Valor": [10, None, 30],
         "Cliques": [1, 2, 3]},
    )
    avisos: list[str] = []
    notas: list[str] = []
    excel = gr.abrir_dados(ficheiro)
    gr.preparar_dados(excel, grafico(), 1, avisos, notas)
    gr.preparar_dados(excel, grafico(titulo="Outro"), 2, avisos, notas)

    assert len(avisos) == len(set(avisos)), f"avisos repetidos: {avisos}"


# --------------------------------------------------- cabecalho fora da linha 1
# Folhas reais costumam ter o titulo do relatorio e a data de exportacao por
# cima da tabela. Sem isto, o titulo virava cabecalho e as colunas ficavam
# todas «Unnamed».


def folha_com_titulo_por_cima(caminho: Path) -> Path:
    from openpyxl import Workbook

    caminho.parent.mkdir(parents=True, exist_ok=True)
    livro = Workbook()
    folha = livro.active
    folha.title = "Vendas"
    folha["A1"] = "Relatório de Campanhas — 1.º trimestre"
    folha["A2"] = "Exportado em 31/03/2026"
    folha.append([])
    folha.append(["Mes", "Valor"])
    for mes, valor in zip(MESES[:3], [10, 20, 35]):
        folha.append([mes, valor])
    livro.save(caminho)
    return caminho


def test_cabecalho_fora_da_primeira_linha_e_encontrado(tmp_path):
    ficheiro = folha_com_titulo_por_cima(tmp_path / "titulo.xlsx")
    serie, n_linhas, avisos, _ = preparar(ficheiro, grafico())

    assert list(serie.index) == MESES[:3]
    assert serie.sum() == 65
    assert n_linhas == 3
    assert any("não começa na primeira linha" in aviso for aviso in avisos)


def test_linha_cabecalho_explicita_nao_gera_aviso(tmp_path):
    ficheiro = folha_com_titulo_por_cima(tmp_path / "titulo.xlsx")
    serie, _, avisos, _ = preparar(ficheiro, grafico(linha_cabecalho=4))

    assert serie.sum() == 65
    assert avisos == []


def test_folha_normal_continua_a_usar_a_primeira_linha(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "normal.xlsx",
        {"Mes": MESES[:3], "Valor": [10, 20, 35]},
    )
    serie, _, avisos, _ = preparar(ficheiro, grafico())
    assert serie.sum() == 65
    assert avisos == []


def test_linha_cabecalho_invalida_e_recusada():
    with pytest.raises(gr.ErroDados, match="linha_cabecalho"):
        gr.validar_grafico(grafico(linha_cabecalho=0), 1)


# ----------------------------------------------------------------- circular


def test_circular_com_negativos_e_erro(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "neg.xlsx",
        {"Mes": ["Janeiro", "Fevereiro"], "Valor": [10, -5]},
    )
    with pytest.raises(gr.ErroDados, match="negativos"):
        preparar(ficheiro, grafico(tipo="circular"))


def test_circular_com_muitas_fatias_avisa(tmp_path):
    categorias = [f"Categoria {i}" for i in range(20)]
    ficheiro = escrever_excel(
        tmp_path / "fatias.xlsx",
        {"Mes": categorias, "Valor": list(range(1, 21))},
    )
    _, _, avisos, _ = preparar(ficheiro, grafico(tipo="circular"))
    assert any("fatias" in aviso for aviso in avisos)


# ------------------------------------------------------------------- CSV


def escrever_csv(caminho: Path, linhas: list[str], separador: str = ";",
                 codificacao: str = "utf-8") -> Path:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    caminho.write_text("\n".join(linhas), encoding=codificacao)
    return caminho


def test_csv_com_ponto_e_virgula(tmp_path):
    """O Excel portugues exporta com «;»."""
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Mes;Valor", "Janeiro;10", "Fevereiro;20", "Marco;30",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert list(serie.index) == ["Janeiro", "Fevereiro", "Marco"]
    assert serie.sum() == 60


def test_csv_com_virgula(tmp_path):
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Mes,Valor", "Janeiro,10", "Fevereiro,20",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie.sum() == 30


def test_csv_em_cp1252_com_acentos(tmp_path):
    """Ficheiros exportados por Windows antigos nao sao UTF-8."""
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Regiao;Valor", "Lisboa;10", "Évora;20", "Bragança;30",
    ], codificacao="cp1252")
    serie, _, _, _ = preparar(ficheiro, grafico(eixo_x="Regiao"))
    assert "Évora" in list(serie.index)
    assert "Bragança" in list(serie.index)


def test_csv_com_metadados_por_cima(tmp_path):
    """Exportacoes do Google Ads trazem linhas de relatorio antes da tabela."""
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Relatorio de campanhas", "Intervalo: ultimo mes", "",
        "Mes;Valor", "Janeiro;10", "Fevereiro;20",
    ])
    serie, _, avisos, _ = preparar(ficheiro, grafico())
    assert serie.sum() == 30
    assert any("não começa na primeira linha" in a for a in avisos)


def test_csv_com_moeda_a_portuguesa(tmp_path):
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Mes;Valor", "Janeiro;1.250,00 €", "Fevereiro;980,50 €",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie["Janeiro"] == pytest.approx(1250.0)
    assert serie["Fevereiro"] == pytest.approx(980.5)


def test_csv_ignora_a_folha_indicada_no_plano(tmp_path):
    ficheiro = escrever_csv(tmp_path / "v.csv", ["Mes;Valor", "Janeiro;10", "Fevereiro;20"])
    avisos: list[str] = []
    notas: list[str] = []
    fonte = gr.abrir_dados(ficheiro)
    gr.preparar_dados(fonte, grafico(folha="Vendas"), 1, avisos, notas)
    assert any("não tem folhas" in n for n in notas)


def test_csv_nao_transforma_texto_em_aviso_bloqueante(tmp_path):
    """Num CSV e tudo texto: se isso fosse aviso, bloqueava sempre e virava ruido."""
    ficheiro = escrever_csv(tmp_path / "v.csv", ["Mes;Valor", "Janeiro;10", "Fevereiro;20"])
    _, _, avisos, notas = preparar(ficheiro, grafico())
    assert avisos == []
    assert any("guardada como texto" in n for n in notas)


def test_csv_com_celulas_perdidas_continua_a_avisar(tmp_path):
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Mes;Valor", "Janeiro;10", "Fevereiro;vinte", "Marco;30",
    ])
    _, _, avisos, _ = preparar(ficheiro, grafico())
    assert any("vinte" in a for a in avisos)


def test_categoria_chamada_NA_nao_desaparece(tmp_path):
    """«NA» e North America, nao «vazio». O default do pandas comia-a."""
    ficheiro = escrever_csv(tmp_path / "na.csv", [
        "Regiao;Valor", "NA;100", "EU;200", "APAC;50",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico(eixo_x="Regiao"))
    assert "NA" in list(serie.index)
    assert serie["NA"] == 100


def test_xlsm_e_aceite(tmp_path):
    caminho = tmp_path / "com_macros.xlsm"
    pd.DataFrame({"Mes": ["Janeiro", "Fevereiro"], "Valor": [10, 20]}).to_excel(
        caminho, sheet_name="Vendas", index=False)
    serie, _, _, _ = preparar(caminho, grafico())
    assert serie.sum() == 30


def test_xls_antigo_da_mensagem_util(tmp_path):
    falso = tmp_path / "antigo.xls"
    falso.write_bytes(b"nao interessa")
    with pytest.raises(gr.ErroDados, match="formato antigo"):
        gr.abrir_dados(falso)


def test_folha_em_falta_com_varias_folhas_e_erro(tmp_path):
    caminho = tmp_path / "duas.xlsx"
    with pd.ExcelWriter(caminho) as escritor:
        pd.DataFrame({"Mes": ["Janeiro"], "Valor": [1]}).to_excel(
            escritor, sheet_name="Vendas", index=False)
        pd.DataFrame({"Mes": ["Janeiro"], "Valor": [2]}).to_excel(
            escritor, sheet_name="Custos", index=False)
    config = grafico()
    config.pop("folha", None)
    with pytest.raises(gr.ErroDados, match="não diz de que folha"):
        preparar(caminho, config)


def test_folha_unica_dispensa_o_campo(tmp_path):
    ficheiro = escrever_excel(tmp_path / "uma.xlsx", {"Mes": ["Janeiro"], "Valor": [7]})
    config = grafico()
    config.pop("folha", None)
    serie, _, _, _ = preparar(ficheiro, config)
    assert serie.sum() == 7


# ---------------------------------------------------------- datas em texto


def test_datas_em_texto_com_convencao_provada(tmp_path):
    """Um dia acima de 12 prova que e dia/mes."""
    ficheiro = escrever_csv(tmp_path / "d.csv", [
        "Data;Valor",
        "31/01/2026;10", "15/02/2026;20", "03/03/2026;30", "20/04/2026;40",
    ])
    serie, _, _, notas = preparar(ficheiro, grafico(eixo_x="Data"))
    assert any("reconhecida como datas" in n for n in notas)
    assert [d.month for d in serie.index] == [1, 2, 3, 4]


def test_datas_ambiguas_ficam_categoricas(tmp_path):
    """Nenhum dia acima de 12: nao da para saber se e dia/mes ou mes/dia."""
    ficheiro = escrever_csv(tmp_path / "d.csv", [
        "Data;Valor", "01/02/2026;10", "03/04/2026;20", "05/06/2026;30",
    ])
    _, _, _, notas = preparar(ficheiro, grafico(eixo_x="Data"))
    assert any("não é possível saber" in n for n in notas)


def test_datas_iso_sao_sempre_aceites(tmp_path):
    ficheiro = escrever_csv(tmp_path / "d.csv", [
        "Data;Valor", "2026-01-05;10", "2026-02-06;20", "2026-03-07;30",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico(eixo_x="Data"))
    assert [d.month for d in serie.index] == [1, 2, 3]


def test_datas_com_convencoes_misturadas_param_com_erro(tmp_path):
    ficheiro = escrever_csv(tmp_path / "d.csv", [
        "Data;Valor", "31/01/2026;10", "01/31/2026;20",
    ])
    with pytest.raises(gr.ErroDados, match="incompatíveis"):
        preparar(ficheiro, grafico(eixo_x="Data"))


def test_separador_detetado_em_varios_formatos():
    assert gr.detetar_separador("a;b;c\n1;2;3") == ";"
    assert gr.detetar_separador("a,b,c\n1,2,3") == ","
    assert gr.detetar_separador("a\tb\tc\n1\t2\t3") == "\t"


# ------------------------------------------------------------ ficheiro e plano


def test_extensao_errada_da_mensagem_clara(tmp_path):
    falso = tmp_path / "antigo.xls"
    falso.write_bytes(b"nao interessa")
    with pytest.raises(gr.ErroDados, match=r"\.xlsx"):
        gr.abrir_dados(falso)


def test_ficheiro_inexistente_da_mensagem_clara(tmp_path):
    with pytest.raises(gr.ErroDados, match="não foi encontrado"):
        gr.abrir_dados(tmp_path / "nao_existe.xlsx")


def test_caminho_com_espacos_e_acentos_funciona(tmp_path):
    pasta = tmp_path / "relatórios de gestão" / "ano corrente"
    ficheiro = escrever_excel(pasta / "vendas anuais.xlsx",
                              {"Mes": MESES[:3], "Valor": [1, 2, 3]})
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie.sum() == 6


def test_nota_com_numeros_e_recusada():
    with pytest.raises(gr.ErroDados, match="números"):
        gr.validar_grafico(grafico(nota="cresceu 20% no verão"), 1)


def test_nota_sem_numeros_passa():
    gr.validar_grafico(grafico(nota="período da campanha de verão"), 1)


def test_campo_que_nao_existe_e_recusado():
    """Apanha planos escritos contra a especificacao antiga."""
    with pytest.raises(gr.ErroDados, match="ordem_categorias"):
        gr.validar_grafico(grafico(ordem_categorias=["Janeiro"]), 1)


def test_tipo_desconhecido_lista_os_disponiveis():
    with pytest.raises(gr.ErroDados) as erro:
        gr.validar_grafico(grafico(tipo="radar"), 1)
    assert "barras" in str(erro.value)


def test_media_sem_eixo_y_e_erro():
    config = grafico(agregacao="media")
    del config["eixo_y"]
    with pytest.raises(gr.ErroDados, match="eixo_y"):
        gr.validar_grafico(config, 1)


def test_contagem_sem_eixo_y_conta_linhas(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "contagem.xlsx",
        {"Mes": ["Janeiro", "Janeiro", "Fevereiro"], "Valor": [1, 2, 3]},
    )
    config = grafico(agregacao="contagem")
    del config["eixo_y"]
    serie, _, _, _ = preparar(ficheiro, config)
    assert serie["Janeiro"] == 2
    assert serie["Fevereiro"] == 1


def test_plano_com_bom_e_aceite(tmp_path):
    """O Bloco de Notas e o PowerShell gravam UTF-8 com BOM. Nao e motivo para recusar."""
    plano = tmp_path / "com_bom.json"
    plano.write_text(
        json.dumps({"titulo_relatorio": "x", "graficos": [grafico()]}, ensure_ascii=False),
        encoding="utf-8-sig",
    )
    assert gr.ler_plano(plano)["titulo_relatorio"] == "x"


def test_plano_com_json_invalido_diz_a_linha(tmp_path):
    plano = tmp_path / "mau.json"
    plano.write_text('{"titulo_relatorio": "x",}', encoding="utf-8")
    with pytest.raises(gr.ErroDados, match="linha"):
        gr.ler_plano(plano)


# ------------------------------------------------------------- estatistica
# Valores conferidos a parte. Um R2 errado nao se ve a olho como se ve um
# grafico partido: tem de ser batido contra outra conta.


def test_regressao_com_reta_perfeita():
    reta = gr.regressao_linear([0.0, 2.0, 4.0, 6.0, 8.0])
    assert reta["declive"] == pytest.approx(2.0)
    assert reta["intercecao"] == pytest.approx(0.0)
    assert reta["r2"] == pytest.approx(1.0)
    assert reta["erro_padrao"] == pytest.approx(0.0)


def test_regressao_com_serie_plana_da_declive_zero():
    reta = gr.regressao_linear([5.0, 5.0, 5.0, 5.0])
    assert reta["declive"] == pytest.approx(0.0)
    assert reta["r2"] == pytest.approx(1.0)


def test_regressao_nao_corre_com_poucos_pontos():
    """Uma reta sobre tres pontos faz um relatorio parecer serio sendo falso."""
    assert gr.regressao_linear([1.0, 2.0, 3.0]) is None


def test_crescimento_medio_duplica_por_periodo():
    # 100 -> 200 -> 400: 100% por periodo
    assert gr.crescimento_medio([100.0, 200.0, 400.0]) == pytest.approx(100.0)


def test_crescimento_medio_recusa_zeros_e_negativos():
    assert gr.crescimento_medio([100.0, 0.0, 400.0]) is None
    assert gr.crescimento_medio([100.0, -50.0, 400.0]) is None


def test_crescimento_medio_recusa_poucos_pontos():
    assert gr.crescimento_medio([100.0, 200.0]) is None


def test_atipicos_encontra_o_valor_disparatado():
    serie = pd.Series([10, 11, 12, 11, 10, 900])
    encontrados = gr.detetar_atipicos(serie)
    assert [v for _, v in encontrados] == [900]


def test_atipicos_nao_correm_com_poucas_categorias():
    assert gr.detetar_atipicos(pd.Series([1, 2, 3, 4])) is None


def test_classificacao_usa_os_limiares_documentados():
    termos = ("baixa", "moderada", "elevada")
    assert gr.classificar(14.9, 15, 35, termos) == "baixa"
    assert gr.classificar(20.0, 15, 35, termos) == "moderada"
    assert gr.classificar(35.1, 15, 35, termos) == "elevada"


# ------------------------------------------- achados em dados publicos reais
# Cada um destes apareceu ao correr a skill sobre dados abertos de terceiros
# (PIB do Banco Mundial, anomalias de temperatura global). Ficam aqui presos
# com ficheiros sinteticos minimos, sem meter dados de ninguem no repositorio.


def test_indice_sazonal_recusa_series_com_negativos():
    """Anomalias de temperatura oscilam a volta de zero: o indice e uma razao
    e explodia, saindo «índice -0,34 — 134% abaixo», um disparate com ar de rigor."""
    rotulos, valores = [], []
    for ano in (2024, 2025):
        for mes in range(1, 13):
            rotulos.append(f"{ano}-{mes:02d}")
            valores.append(-0.5 if mes % 2 else 0.4)
    resultado = gr.indice_sazonal(pd.Series(valores, index=rotulos))
    assert resultado["impossivel"] == "negativos"


def test_seccao_sazonalidade_explica_a_recusa_por_negativos():
    rotulos, valores = [], []
    for ano in (2024, 2025):
        for mes in range(1, 13):
            rotulos.append(f"{ano}-{mes:02d}")
            valores.append(-0.5 if mes % 2 else 0.4)
    blocos = gr.montar_analise(pd.Series(valores, index=rotulos), grafico(), 24, True, None)
    assert "positivos" in texto_de(blocos, "Sazonalidade")


def test_declive_pequeno_nao_e_arredondado_para_zero():
    """Um declive de 0,0005 aparecia como «+0», o que nao diz nada."""
    valores = [i * 0.0005 for i in range(12)]
    serie = pd.Series(valores, index=[f"P{i}" for i in range(12)])
    blocos = gr.montar_analise(serie, grafico(), 12, True, None)
    assert "+0 por período" not in texto_de(blocos, "Tendência")
    assert "0,0005" in texto_de(blocos, "Tendência")


def test_previsao_de_valores_pequenos_nao_colapsa_em_inteiros():
    """A previsao arredondada a inteiros dava «entre 0 e 1» em dados pequenos."""
    serie = pd.Series([i * 0.01 for i in range(12)], index=[f"P{i}" for i in range(12)])
    _, texto = gr.secao_previsao(serie, 3)
    assert "entre 0 e 1" not in texto


def test_anos_anteriores_a_1900_sao_reconhecidos():
    """A serie de temperatura comeca em 1850 e ficava como categoria."""
    assert gr.e_ano(1850) is True
    assert gr.extrair_ano("1850-01") == 1850
    assert gr.parece_temporal(pd.Index(["1850-01", "1850-02", "1850-03"])) is True


def test_formato_ano_mes_e_reconhecido_sem_ajuda():
    """As exportacoes de plataformas trazem «2023-01» e nenhuma dica no plano."""
    assert gr.parece_temporal(pd.Index(["2023-01", "2023-02", "2023-03"])) is True


def test_barras_com_categorias_a_mais_avisam(tmp_path):
    """262 paises num grafico de barras sao ilegiveis, e nada avisava."""
    categorias = [f"Pais {i}" for i in range(60)]
    ficheiro = escrever_excel(tmp_path / "muitos.xlsx",
                              {"Mes": categorias, "Valor": list(range(60))})
    _, _, avisos, _ = preparar(ficheiro, grafico())
    assert any("categorias" in a and "ilegíveis" in a for a in avisos)


def test_linhas_com_muitos_pontos_nao_avisam(tmp_path):
    """Numa linha, 60 pontos leem-se bem; o aviso das barras seria falso positivo."""
    categorias = [f"P{i}" for i in range(60)]
    ficheiro = escrever_excel(tmp_path / "linha.xlsx",
                              {"Mes": categorias, "Valor": list(range(60))})
    _, _, avisos, _ = preparar(ficheiro, grafico(tipo="linhas"))
    assert not any("ilegíveis" in a for a in avisos)


def test_tabela_truncada_avisa_antes_de_gerar(tmp_path):
    """O README dizia que bloqueava; o codigo so o registava depois de gerar."""
    categorias = [f"C{i}" for i in range(50)]
    ficheiro = escrever_excel(tmp_path / "tabela.xlsx",
                              {"Mes": categorias, "Valor": list(range(50))})
    _, _, avisos, _ = preparar(ficheiro, grafico(tipo="linhas", tabela_dados=True))
    assert any("ficariam de fora" in a for a in avisos)


# ------------------------------------------------------ sugerir graficos


def sugerir(caminho: Path) -> str:
    return correr_cli("--dados", str(caminho), "--sugerir").stdout


def ficheiro_tipico(tmp_path: Path) -> Path:
    return escrever_excel(tmp_path / "s.xlsx", {
        "Periodo": [f"2024-{m:02d}" for m in range(1, 7)] * 2,
        "Canal": ["Online"] * 6 + ["Loja"] * 6,
        "Valor": list(range(100, 220, 10)),
        "Encomendas": list(range(10, 22)),
    })


def test_sugere_os_graficos_obvios(tmp_path):
    saida = sugerir(ficheiro_tipico(tmp_path))

    assert "Evolução de Valor ao longo de Periodo" in saida
    assert "séries" in saida            # comparar canais no mesmo grafico
    assert "circular" in saida          # 2 canais, cabe num circular
    assert "dispersão" in saida         # Valor e Encomendas


def test_cada_sugestao_diz_porque(tmp_path):
    """Sem a razao sao palpites; com ela, da para discordar da regra."""
    saida = sugerir(ficheiro_tipico(tmp_path))
    assert saida.count("porquê:") >= 3


def test_coluna_de_datas_nao_e_confundida_com_identificador(tmp_path):
    """Numa serie diaria cada data e unica: a regra do identificador matava
    exatamente a coluna mais valiosa. Apanhado com dados reais."""
    ficheiro = escrever_csv(tmp_path / "d.csv", ["data;valor"] + [
        f"{dia:02d}-01-2025;{dia}" for dia in range(1, 21)
    ])
    saida = sugerir(ficheiro)
    assert "data — linha do tempo" in saida
    assert "Evolução" in saida


def test_identificadores_ficam_de_fora(tmp_path):
    ficheiro = escrever_excel(tmp_path / "i.xlsx", {
        "id_cliente": list(range(1, 11)),
        "Regiao": ["Norte", "Sul"] * 5,
        "Valor": list(range(10, 20)),
    })
    saida = sugerir(ficheiro)
    assert "deixei de fora" in saida
    assert "id_cliente" in saida.split("Sugestões")[0]


def test_nao_cruza_um_total_com_uma_parte_dele():
    """«confirmados» contra «confirmados_arsnorte» da correlacao por construcao."""
    numeros = [{"nome": "confirmados"}, {"nome": "confirmados_arsnorte"}]
    assert gr.escolher_par_de_numeros(numeros) is None

    numeros.append({"nome": "recuperados"})
    assert gr.escolher_par_de_numeros(numeros) == ("confirmados", "recuperados")


def test_lista_de_colunas_nao_enche_o_ecra(tmp_path):
    dados = {"Mes": MESES}
    for i in range(40):
        dados[f"medida_{i}"] = list(range(12))
    ficheiro = escrever_excel(tmp_path / "muitas.xlsx", dados)
    saida = sugerir(ficheiro)
    assert "e mais" in saida
    assert saida.count("— número") <= gr.MAX_COLUNAS_LISTADAS


def test_ficheiro_sem_nada_para_desenhar_diz_isso(tmp_path):
    ficheiro = escrever_excel(tmp_path / "vazio.xlsx", {
        "notas": [f"comentário livre número {i}" for i in range(60)],
    })
    saida = sugerir(ficheiro)
    assert "Não encontrei colunas" in saida


def test_sugerir_nao_precisa_de_plano(tmp_path):
    resultado = correr_cli("--dados", str(ficheiro_tipico(tmp_path)), "--sugerir")
    assert resultado.returncode == 0


def test_sem_plano_e_sem_sugerir_explica_o_que_falta(tmp_path):
    resultado = correr_cli("--dados", str(ficheiro_tipico(tmp_path)))
    assert resultado.returncode != 0
    assert "--sugerir" in resultado.stderr


# --------------------------------------------- nomes de coluna legiveis


def test_nome_legivel_arruma_o_que_precisa():
    assert gr.nome_legivel("doses_novas") == "Doses novas"
    assert gr.nome_legivel("valor") == "Valor"
    assert gr.nome_legivel("confirmados_arsnorte") == "Confirmados arsnorte"


def test_nome_legivel_nao_estraga_o_que_ja_esta_bem():
    """Se alguem pos maiusculas, foi de proposito."""
    assert gr.nome_legivel("Valor") == "Valor"
    assert gr.nome_legivel("IVA") == "IVA"
    assert gr.nome_legivel("Vendas Brutas") == "Vendas Brutas"
    assert gr.nome_legivel("Periodo") == "Periodo"


def test_relatorio_usa_o_nome_bonito_e_o_aviso_o_real(tmp_path):
    """A distincao que importa: o texto e para ler, o aviso e para procurar."""
    ficheiro = escrever_csv(tmp_path / "n.csv", [
        "mes;doses_novas", "Janeiro;10", "Fevereiro;vinte", "Marco;30",
    ])
    config = grafico(eixo_x="mes", eixo_y="doses_novas")
    serie, _, avisos, _ = preparar(ficheiro, config)

    frase = gr.frase_descritiva(serie, config, 2)
    assert "Doses novas" in frase, "o relatório devia usar o nome bonito"
    assert "doses_novas" not in frase

    assert any("doses_novas" in a for a in avisos), \
        "o aviso devia usar o nome real, para se encontrar a coluna no Excel"


# ------------------------------------------- amostra antes de gerar


def test_verificar_mostra_os_numeros(tmp_path):
    ficheiro = escrever_excel(tmp_path / "a.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    plano = escrever_plano(tmp_path / "p.json", [grafico()])
    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--verificar")

    assert "Vou fazer 1 gráfico" in resultado.stdout
    assert "Janeiro 10" in resultado.stdout
    assert "Total: 60" in resultado.stdout


def test_amostra_corta_as_cinco_primeiras(tmp_path):
    ficheiro = escrever_excel(tmp_path / "a.xlsx",
                              {"Mes": MESES, "Valor": list(range(1, 13))})
    plano = escrever_plano(tmp_path / "p.json", [grafico()])
    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--verificar")

    assert "(mais 7)" in resultado.stdout
    assert "Dezembro" not in resultado.stdout.split("Total:")[0]


def test_amostra_com_series_mostra_as_series(tmp_path):
    ficheiro = escrever_excel(tmp_path / "s.xlsx", {
        "Mes": ["Janeiro", "Janeiro", "Fevereiro", "Fevereiro"],
        "Canal": ["Online", "Loja"] * 2,
        "Valor": [100, 40, 200, 50],
    })
    plano = escrever_plano(tmp_path / "p.json", [grafico(serie="Canal")])
    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--verificar")

    assert "2 séries" in resultado.stdout
    assert "Online 300" in resultado.stdout
    assert "Loja 90" in resultado.stdout


def test_amostra_da_dispersao_mostra_intervalos(tmp_path):
    ficheiro = escrever_excel(tmp_path / "d.xlsx",
                              {"a": [1, 2, 3, 4], "b": [10, 20, 30, 40]})
    plano = escrever_plano(tmp_path / "p.json", [
        {"tipo": "dispersao", "eixo_x": "a", "eixo_y": "b", "titulo": "R"}])
    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--verificar")

    assert "4 pontos" in resultado.stdout
    assert "de 1 a 4" in resultado.stdout
    assert "de 10 a 40" in resultado.stdout


def test_media_mostra_media_e_nao_total(tmp_path):
    ficheiro = escrever_excel(tmp_path / "m.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    plano = escrever_plano(tmp_path / "p.json", [grafico(agregacao="media")])
    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--verificar")

    assert "Média geral: 20" in resultado.stdout
    assert "Total:" not in resultado.stdout


# ------------------------------------------------------------ filtro


def ficheiro_com_anos(tmp_path: Path) -> Path:
    return escrever_excel(tmp_path / "anos.xlsx", {
        "Ano": [2023, 2023, 2024, 2025, 2025, 2025],
        "Mes": ["Janeiro", "Fevereiro"] * 3,
        "Valor": [10, 20, 30, 40, 50, 60],
    })


def test_filtro_igual_a_um_valor(tmp_path):
    serie, n_linhas, _, notas = preparar(
        ficheiro_com_anos(tmp_path),
        grafico(filtro={"coluna": "Ano", "igual_a": 2025}))
    assert n_linhas == 3
    assert serie.sum() == 150


def test_filtro_regista_sempre_o_que_ficou_de_fora(tmp_path):
    """Filtrar sem dizer e esconder dados."""
    _, _, _, notas = preparar(
        ficheiro_com_anos(tmp_path),
        grafico(filtro={"coluna": "Ano", "igual_a": 2025}))
    texto = " ".join(notas)
    assert "filtro aplicado" in texto
    assert "3 de 6 linhas" in texto
    assert "3 ficaram de fora" in texto


def test_filtro_com_lista_de_valores(tmp_path):
    serie, n_linhas, _, _ = preparar(
        ficheiro_com_anos(tmp_path),
        grafico(filtro={"coluna": "Ano", "igual_a": [2023, 2024]}))
    assert n_linhas == 3
    assert serie.sum() == 60


def test_filtro_por_intervalo(tmp_path):
    serie, _, _, _ = preparar(
        ficheiro_com_anos(tmp_path),
        grafico(filtro={"coluna": "Valor", "de": 30, "ate": 50}))
    assert serie.sum() == 120  # 30 + 40 + 50


def test_filtro_que_nao_deixa_nada_e_erro(tmp_path):
    with pytest.raises(gr.ErroDados, match="não deixou nenhuma linha"):
        preparar(ficheiro_com_anos(tmp_path),
                 grafico(filtro={"coluna": "Ano", "igual_a": 1999}))


def test_filtro_com_coluna_inexistente_e_erro(tmp_path):
    with pytest.raises(gr.ErroDados, match="Regiao"):
        preparar(ficheiro_com_anos(tmp_path),
                 grafico(filtro={"coluna": "Regiao", "igual_a": "Norte"}))


def test_filtro_mal_formado_e_recusado():
    with pytest.raises(gr.ErroDados, match="objeto"):
        gr.validar_grafico(grafico(filtro="Ano=2025"), 1)
    with pytest.raises(gr.ErroDados, match="coluna"):
        gr.validar_grafico(grafico(filtro={"igual_a": 2025}), 1)
    with pytest.raises(gr.ErroDados, match="não diz o que filtrar"):
        gr.validar_grafico(grafico(filtro={"coluna": "Ano"}), 1)


# ------------------------------------------------- tipos de grafico novos


def test_barras_horizontais_e_area_funcionam(tmp_path):
    ficheiro = escrever_excel(tmp_path / "t.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    for tipo in ("barras_horizontais", "area"):
        serie, _, _, _ = preparar(ficheiro, grafico(tipo=tipo))
        assert serie.sum() == 60, tipo


def test_dispersao_cruza_duas_colunas(tmp_path):
    ficheiro = escrever_excel(tmp_path / "d.xlsx", {
        "Investimento": [100, 200, 300, 400],
        "Conversoes": [10, 19, 31, 40],
    })
    config = {"tipo": "dispersao", "eixo_x": "Investimento",
              "eixo_y": "Conversoes", "titulo": "Relação"}
    serie, n_linhas, _, _ = preparar(ficheiro, config)

    assert n_linhas == 4  # um ponto por linha, sem agrupar
    assert list(serie.values) == [10, 19, 31, 40]


def test_dispersao_mede_a_correlacao_e_avisa_que_nao_e_causa(tmp_path):
    ficheiro = escrever_excel(tmp_path / "d.xlsx", {
        "Investimento": [100, 200, 300, 400],
        "Conversoes": [10, 20, 30, 40],
    })
    config = {"tipo": "dispersao", "eixo_x": "Investimento",
              "eixo_y": "Conversoes", "titulo": "Relação"}
    serie, _, _, _ = preparar(ficheiro, config)
    blocos = gr.montar_analise(serie, config, 4, False, None)

    assert "1" in texto_de(blocos, "Relação")  # correlacao perfeita
    assert "forte" in texto_de(blocos, "Relação")
    assert "não é causa" in texto_de(blocos, "O que isto não diz")


def test_correlacao_conferida_a_mao():
    assert gr.correlacao([1, 2, 3, 4], [2, 4, 6, 8]) == pytest.approx(1.0)
    assert gr.correlacao([1, 2, 3, 4], [8, 6, 4, 2]) == pytest.approx(-1.0)
    assert gr.correlacao([1, 1, 1], [1, 2, 3]) is None


def test_dispersao_com_agregacao_e_recusada():
    with pytest.raises(gr.ErroDados, match="Tira a «agregacao»"):
        gr.validar_grafico(grafico(tipo="dispersao"), 1)


def test_dispersao_precisa_de_eixo_y():
    with pytest.raises(gr.ErroDados, match="duas colunas"):
        gr.validar_grafico({"tipo": "dispersao", "eixo_x": "a", "titulo": "t"}, 1)


def test_dispersao_com_poucos_pontos_e_erro(tmp_path):
    ficheiro = escrever_excel(tmp_path / "d.xlsx",
                              {"a": [1, 2], "b": [3, 4]})
    config = {"tipo": "dispersao", "eixo_x": "a", "eixo_y": "b", "titulo": "t"}
    with pytest.raises(gr.ErroDados, match="pelo menos 3"):
        preparar(ficheiro, config)


def test_agregacao_continua_obrigatoria_nos_outros_tipos():
    config = grafico()
    config.pop("agregacao")
    with pytest.raises(gr.ErroDados, match="Falta «agregacao»"):
        gr.validar_grafico(config, 1)


# ------------------------------------------------ capa, indice e paginas


def gerar_para_ler(tmp_path: Path, graficos: list[dict], **plano_extra):
    """Gera um documento e devolve-o aberto, para inspecionar."""
    from docx import Document

    ficheiro = escrever_excel(tmp_path / "d.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    plano = tmp_path / "p.json"
    corpo = {"titulo_relatorio": "Relatório de teste", "graficos": graficos}
    corpo.update(plano_extra)
    plano.write_text(json.dumps(corpo, ensure_ascii=False), encoding="utf-8")
    saida = tmp_path / "r.docx"
    correr_cli("--dados", str(ficheiro), "--plano", str(plano), "--saida", str(saida))
    return Document(str(saida))


def test_capa_traz_titulo_data_e_origem(tmp_path):
    import datetime

    documento = gerar_para_ler(tmp_path, [grafico()])
    texto = "\n".join(p.text for p in documento.paragraphs)

    assert "Relatório de teste" in texto
    assert datetime.date.today().strftime("%d/%m/%Y") in texto
    assert "d.xlsx" in texto


def test_indice_lista_todos_os_graficos(tmp_path):
    documento = gerar_para_ler(tmp_path, [
        grafico(titulo="Primeiro"), grafico(titulo="Segundo"),
    ])
    texto = "\n".join(p.text for p in documento.paragraphs)

    assert "Índice" in texto
    assert "1.  Primeiro" in texto
    assert "2.  Segundo" in texto


def test_rodape_tem_numero_de_pagina(tmp_path):
    documento = gerar_para_ler(tmp_path, [grafico()])
    seccao = documento.sections[0]

    assert seccao.different_first_page_header_footer is True, "a capa não devia ser numerada"
    assert "PAGE" in seccao.footer.paragraphs[0]._p.xml


def test_capa_nao_conta_como_grafico(tmp_path):
    """A capa e o indice nao podem trazer imagens nem mexer nos numeros."""
    documento = gerar_para_ler(tmp_path, [grafico()])
    assert len(documento.inline_shapes) == 1
    texto = "\n".join(p.text for p in documento.paragraphs)
    assert "O total é 60." in texto


# ------------------------------------- achados a usar a skill a serio
# Apareceram ao seguir o SKILL.md a risca sobre dados publicos reais
# (casos de covid em Portugal, 782 dias, 93 colunas, formato largo).


def test_cabecalho_e_encontrado_com_vazios_na_primeira_linha_de_dados(tmp_path):
    """A regra antiga comparava com a linha seguinte, e bastava um vazio nessa
    linha para o cabecalho verdadeiro ser rejeitado. Em dados reais e banal."""
    ficheiro = escrever_csv(tmp_path / "v.csv", [
        "Mes;Valor;Extra",
        "Janeiro;10;",      # a primeira linha de dados tem uma celula vazia
        "Fevereiro;20;x",
        "Marco;30;y",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert list(serie.index) == ["Janeiro", "Fevereiro", "Marco"]
    assert serie.sum() == 60


def test_lista_de_colunas_da_varias_series(tmp_path):
    """Formato largo: uma coluna por regiao, como em muitas exportacoes."""
    ficheiro = escrever_excel(tmp_path / "largo.xlsx", {
        "Mes": MESES[:3], "norte": [10, 20, 30], "sul": [5, 6, 7],
    })
    config = grafico(serie=["norte", "sul"])
    config.pop("eixo_y")
    series, _, _, _ = preparar_series(ficheiro, config)

    # os nomes das series vem de nomes de coluna: entram bonitos na legenda
    assert list(series) == ["Norte", "Sul"]
    assert series["Norte"].sum() == 60
    assert series["Sul"].sum() == 18


def test_lista_de_colunas_com_eixo_y_e_recusada():
    config = grafico(serie=["a", "b"])
    with pytest.raises(gr.ErroDados, match="tira o «eixo_y»"):
        gr.validar_grafico(config, 1)


def test_lista_de_colunas_precisa_de_duas():
    config = grafico(serie=["so_uma"])
    config.pop("eixo_y")
    with pytest.raises(gr.ErroDados, match="pelo menos dois"):
        gr.validar_grafico(config, 1)


def test_serie_acumulada_e_apanhada(tmp_path):
    """Somar um acumulado da um numero sem significado, e nada o denunciava."""
    valores = list(range(1, 31))  # sempre a subir, 30 periodos
    acumulado = [sum(valores[:i + 1]) for i in range(len(valores))]
    ficheiro = escrever_excel(tmp_path / "acum.xlsx", {
        "Mes": [f"2024-{i:02d}" for i in range(1, 31)], "Valor": acumulado,
    })
    _, _, avisos, _ = preparar(ficheiro, grafico())
    assert any("acumulado" in a for a in avisos)


def test_serie_curta_a_crescer_nao_e_confundida_com_acumulado(tmp_path):
    """12 meses sempre a subir sao dados legitimos, nao um acumulado."""
    ficheiro = escrever_excel(tmp_path / "cresce.xlsx",
                              {"Mes": MESES, "Valor": list(range(1, 13))})
    _, _, avisos, _ = preparar(ficheiro, grafico())
    assert not any("acumulado" in a for a in avisos)


def test_serie_sem_descidas_e_monotona_nao_decrescente():
    """Com zeros pelo meio e zero descidas, «não monótona» estava errado."""
    serie = pd.Series([10, 10, 20, 20, 30], index=[f"P{i}" for i in range(5)])
    blocos = gr.montar_analise(serie, grafico(), 5, True, None)
    texto = texto_de(blocos, "Evolução")
    assert "não decrescente" in texto
    assert "não monótona" not in texto


def test_lista_de_colunas_em_falta_lista_as_disponiveis(tmp_path):
    ficheiro = escrever_excel(tmp_path / "largo.xlsx", {
        "Mes": MESES[:3], "norte": [10, 20, 30], "sul": [5, 6, 7],
    })
    config = grafico(serie=["norte", "poente"])
    config.pop("eixo_y")
    with pytest.raises(gr.ErroDados) as erro:
        preparar(ficheiro, config)
    assert "sul" in str(erro.value)


def test_lista_de_colunas_nao_despeja_93_nomes():
    muitas = [f"coluna_{i}" for i in range(93)]
    texto = gr.listar_colunas(muitas)
    assert "e mais 81" in texto
    assert "coluna_92" not in texto


def test_preposicao_certa_na_mensagem():
    """Saia «não existe do ficheiro»."""
    assert gr.onde_fica(gr.FonteCSV.NOME_UNICO, "em") == "no ficheiro"
    assert gr.onde_fica("Vendas", "em") == "na folha «Vendas»"
    assert gr.onde_fica("Vendas") == "da folha «Vendas»"


# --------------------------------------------------------------- series
# Varias series no mesmo grafico: comparar canais ao longo do tempo.


def ficheiro_com_canais(tmp_path: Path) -> Path:
    return escrever_excel(tmp_path / "canais.xlsx", {
        "Mes": ["Janeiro", "Janeiro", "Fevereiro", "Fevereiro", "Marco", "Marco"],
        "Canal": ["Online", "Loja", "Online", "Loja", "Online", "Loja"],
        "Valor": [100, 40, 200, 50, 300, 60],
    })


def test_sem_campo_serie_ha_uma_serie_so(tmp_path):
    """O caso de sempre tem de continuar exatamente igual."""
    ficheiro = escrever_excel(tmp_path / "s.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    series, _, _, _ = preparar_series(ficheiro, grafico())
    assert list(series) == [""]
    assert series[""].sum() == 60


def test_pivo_separa_as_series(tmp_path):
    series, _, _, _ = preparar_series(ficheiro_com_canais(tmp_path),
                                      grafico(serie="Canal"))
    assert list(series) == ["Online", "Loja"]  # ordem de aparicao no ficheiro
    assert series["Online"].sum() == 600
    assert series["Loja"].sum() == 150


def test_conjunto_nao_e_a_soma_das_series_com_media(tmp_path):
    """Com «media», somar as medias de cada serie daria um numero errado."""
    conjunto, _, _, _ = preparar(ficheiro_com_canais(tmp_path),
                                 grafico(serie="Canal", agregacao="media"))
    # media de todas as linhas de Janeiro: (100+40)/2 = 70
    assert conjunto["Janeiro"] == pytest.approx(70.0)


def test_circular_com_series_e_erro(tmp_path):
    with pytest.raises(gr.ErroDados, match="circular"):
        preparar(ficheiro_com_canais(tmp_path),
                 grafico(tipo="circular", serie="Canal"))


def test_serie_igual_ao_eixo_x_e_erro(tmp_path):
    with pytest.raises(gr.ErroDados, match="ao mesmo tempo"):
        preparar(ficheiro_com_canais(tmp_path), grafico(serie="Mes"))


def test_coluna_de_serie_inexistente_lista_as_que_existem(tmp_path):
    with pytest.raises(gr.ErroDados) as erro:
        preparar(ficheiro_com_canais(tmp_path), grafico(serie="Regiao"))
    assert "Canal" in str(erro.value)


def test_series_a_mais_avisam(tmp_path):
    ficheiro = escrever_excel(tmp_path / "muitas.xlsx", {
        "Mes": ["Janeiro"] * 8,
        "Canal": [f"C{i}" for i in range(8)],
        "Valor": list(range(8)),
    })
    _, _, avisos, _ = preparar_series(ficheiro, grafico(serie="Canal"))
    assert any("séries no mesmo gráfico" in a for a in avisos)


def test_serie_que_nao_cobre_todas_as_categorias_fica_alinhada(tmp_path):
    """Um canal que so existiu a partir de certa altura nao pode desalinhar
    as linhas dos rotulos. A falta entra como buraco, nunca como zero."""
    ficheiro = escrever_excel(tmp_path / "parcial.xlsx", {
        "Mes": ["Janeiro", "Fevereiro", "Fevereiro", "Marco"],
        "Canal": ["Online", "Online", "Novo", "Novo"],
        "Valor": [10, 20, 5, 7],
    })
    series, _, _, _ = preparar_series(ficheiro, grafico(serie="Canal"))
    conjunto, _, _, _ = preparar(ficheiro, grafico(serie="Canal"))

    grelha = gr.alinhar(series["Novo"], conjunto.index)
    assert len(grelha) == len(conjunto)
    assert grelha[0] != grelha[0], "Janeiro devia ser buraco (NaN), nunca zero"
    assert grelha[1] == pytest.approx(5.0)


def test_comparacao_entre_series_traz_pesos(tmp_path):
    series, temporal, _, _ = preparar_series(ficheiro_com_canais(tmp_path),
                                             grafico(serie="Canal"))
    blocos = gr.secao_comparacao_series(series, grafico(serie="Canal"), temporal)
    texto = dict(blocos)["Comparação entre séries"]
    assert "2 séries" in texto
    assert "80%" in texto  # 600 de 750


def test_meta_por_serie_conta_as_que_atingiram(tmp_path):
    series, _, _, _ = preparar_series(ficheiro_com_canais(tmp_path),
                                      grafico(serie="Canal"))
    config = grafico(serie="Canal", meta={"valor": 200, "ambito": "serie"})
    _, texto = gr.secao_meta_por_serie(series, config)
    assert "1 de 2" in texto
    assert "Loja" in texto


def test_meta_ambito_serie_sem_series_e_erro():
    with pytest.raises(gr.ErroDados, match="não tem séries"):
        gr.validar_grafico(grafico(meta={"valor": 10, "ambito": "serie"}), 1)


def test_meta_ambito_categoria_com_series_e_ambiguo():
    with pytest.raises(gr.ErroDados, match="ambíguo"):
        gr.validar_grafico(
            grafico(serie="Canal", meta={"valor": 10, "ambito": "categoria"}), 1)


def test_analise_detalhada_repete_por_serie(tmp_path):
    from docx import Document

    ficheiro = escrever_excel(tmp_path / "d.xlsx", {
        "Mes": MESES[:3] * 2,
        "Canal": ["Online"] * 3 + ["Loja"] * 3,
        "Valor": [10, 20, 30, 5, 6, 7],
    })
    for profundidade, esperado in (("completa", False), ("detalhada", True)):
        plano = tmp_path / f"{profundidade}.json"
        plano.write_text(json.dumps({
            "titulo_relatorio": "T", "analise": profundidade,
            "graficos": [grafico(serie="Canal")],
        }, ensure_ascii=False), encoding="utf-8")
        saida = tmp_path / f"{profundidade}.docx"
        correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                   "--saida", str(saida))
        texto = "\n".join(p.text for p in Document(str(saida)).paragraphs)
        assert "Comparação entre séries" in texto
        assert ("Análise — Online" in texto) is esperado


# ------------------------------------------------ achados na revisao critica
# Quatro defeitos encontrados a reler o codigo com desconfianca, depois de
# escrito. Nenhum deles tinha teste.


def test_csv_com_cabecalhos_repetidos_nao_rebenta(tmp_path):
    """Exportacoes trazem colunas com o mesmo nome; df[coluna] devolvia um
    DataFrame em vez de uma coluna e saia um traceback."""
    ficheiro = escrever_csv(tmp_path / "dup.csv", [
        "Mes;Valor;Valor", "Janeiro;10;99", "Fevereiro;20;99",
    ])
    serie, _, _, _ = preparar(ficheiro, grafico())
    assert serie.sum() == 30  # ficou com a primeira coluna «Valor»


def test_cabecalhos_repetidos_sao_desambiguados():
    assert gr.nomes_unicos(["a", "b", "a", "a", "b"]) == ["a", "b", "a.1", "a.2", "b.1"]


def test_meta_com_analise_curta_e_recusada(tmp_path):
    """Antes eram ignoradas em silencio: pedias avaliacao e nao recebias nada."""
    plano = tmp_path / "p.json"
    plano.write_text(json.dumps({
        "titulo_relatorio": "T", "analise": "curta",
        "graficos": [grafico(meta={"valor": 100})],
    }, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(gr.ErroDados, match="curta"):
        gr.ler_plano(plano)


def test_previsao_com_analise_curta_e_recusada(tmp_path):
    plano = tmp_path / "p.json"
    plano.write_text(json.dumps({
        "titulo_relatorio": "T", "analise": "curta",
        "graficos": [grafico(previsao=3)],
    }, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(gr.ErroDados, match="curta"):
        gr.ler_plano(plano)


def test_centimos_nao_desaparecem():
    """O numero impresso nao pode contradizer o veredicto ao lado."""
    assert gr.formatar_com_precisao(7745.75) == "7.745,75"
    assert gr.formatar_com_precisao(0.4) == "0,40"
    assert gr.formatar_com_precisao(1214) == "1.214"  # inteiro fica inteiro
    assert gr.formatar_com_precisao(0.000561) == "0,000561"


def test_diferenca_pequena_nao_e_impressa_como_zero():
    serie = pd.Series([100.0, 100.4], index=["a", "b"])
    config = grafico(meta={"valor": 200, "ambito": "total"})
    _, texto = gr.secao_meta(serie, config)
    assert "0,40 acima da meta" in texto


def test_meta_negativa_e_recusada():
    with pytest.raises(gr.ErroDados, match="negativa"):
        gr.validar_grafico(grafico(meta={"valor": -500}), 1)


# -------------------------------------------------------------------- meta
# A unica seccao que avalia desempenho, e so porque a referencia vem de fora.


def serie_simples() -> pd.Series:
    return pd.Series([100.0, 200.0, 300.0], index=["Janeiro", "Fevereiro", "Marco"])


def test_meta_total_abaixo():
    config = grafico(meta={"valor": 800, "ambito": "total"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "75%" in texto
    assert "abaixo da meta" in texto
    assert "200" in texto  # o que faltou


def test_meta_total_acima():
    config = grafico(meta={"valor": 500, "ambito": "total"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "120%" in texto
    assert "acima da meta" in texto


def test_meta_total_exata():
    config = grafico(meta={"valor": 600, "ambito": "total"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "exatamente na meta" in texto


def test_meta_com_media_compara_a_media():
    config = grafico(agregacao="media", meta={"valor": 200, "ambito": "total"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "média das categorias" in texto.lower()
    assert "100%" in texto


def test_meta_por_categoria_conta_as_que_atingiram():
    config = grafico(meta={"valor": 200, "ambito": "categoria"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "2 de 3" in texto
    assert "Janeiro" in texto  # a que ficou abaixo
    assert "défice total é 100" in texto


def test_meta_por_categoria_quando_todas_atingem():
    config = grafico(meta={"valor": 50, "ambito": "categoria"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "Nenhuma categoria ficou abaixo" in texto


def test_meta_ambito_por_omissao_e_o_total():
    config = grafico(meta={"valor": 600})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "O total" in texto


def test_meta_zero_nao_divide_por_zero():
    config = grafico(meta={"valor": 0, "ambito": "total"})
    _, texto = gr.secao_meta(serie_simples(), config)
    assert "meta de zero" in texto


def test_meta_entra_na_analise():
    config = grafico(meta={"valor": 600, "ambito": "total"})
    blocos = gr.montar_analise(serie_simples(), config, 3, False, None)
    assert "Meta" in etiquetas(blocos)


def test_sem_meta_nao_ha_seccao_meta():
    blocos = gr.montar_analise(serie_simples(), grafico(), 3, False, None)
    assert "Meta" not in etiquetas(blocos)


def test_meta_continua_sem_juizos_de_valor():
    """Com meta ha avaliacao, mas continua a ser conta e nao opiniao."""
    for alvo in (100, 600, 5000):
        config = grafico(meta={"valor": alvo, "ambito": "total"})
        _, texto = gr.secao_meta(serie_simples(), config)
        baixo = texto.lower()
        for proibida in ("bom", "mau", "fraco", "excelente", "preocupante", "falhou"):
            assert proibida not in baixo, f"apareceu «{proibida}» com meta {alvo}"


def test_meta_mal_formada_e_recusada():
    with pytest.raises(gr.ErroDados, match="objeto"):
        gr.validar_grafico(grafico(meta=700000), 1)
    with pytest.raises(gr.ErroDados, match="valor"):
        gr.validar_grafico(grafico(meta={"ambito": "total"}), 1)
    with pytest.raises(gr.ErroDados, match="âmbito"):
        gr.validar_grafico(grafico(meta={"valor": 10, "ambito": "trimestre"}), 1)
    with pytest.raises(gr.ErroDados, match="não existem"):
        gr.validar_grafico(grafico(meta={"valor": 10, "objetivo": "x"}), 1)


# ---------------------------------------------------------------- previsao


def serie_reta(n: int, declive: float = 2.0) -> pd.Series:
    """Serie perfeitamente linear, com rotulos que nao ativam a sazonalidade."""
    return pd.Series([declive * i for i in range(n)],
                     index=[f"P{i}" for i in range(n)])


def test_previsao_acerta_em_cheio_numa_reta_sem_ruido():
    """Sanidade matematica: se falhar aqui, esta errada em todo o lado."""
    resultado = gr.prever(serie_reta(10), 3)
    previstos = [p["centro"] for p in resultado["previsoes"]]

    assert previstos == pytest.approx([20.0, 22.0, 24.0])
    for p in resultado["previsoes"]:
        assert p["superior"] - p["inferior"] == pytest.approx(0.0, abs=1e-6)


def test_previsao_traz_sempre_intervalo():
    serie = pd.Series([10, 13, 11, 18, 21, 19, 26, 30, 28, 35],
                      index=[f"P{i}" for i in range(10)])
    resultado = gr.prever(serie, 2)
    for p in resultado["previsoes"]:
        assert p["inferior"] < p["centro"] < p["superior"]


def test_previsao_recusa_serie_curta():
    assert gr.prever(serie_reta(5), 2) is None


def test_previsao_recusa_ajuste_fraco():
    """Uma serie sem forma nao se extrapola."""
    serie = pd.Series([10, 90, 20, 85, 15, 95, 25, 80, 12, 88],
                      index=[f"P{i}" for i in range(10)])
    resultado = gr.prever(serie, 2)
    assert resultado["recusa"] == "ajuste"
    assert resultado["r2"] < gr.R2_MINIMO_PREVISAO


def test_horizonte_limitado_a_um_terco_da_serie():
    assert gr.horizonte_permitido(9) == 3
    assert gr.horizonte_permitido(36) == 12
    assert gr.horizonte_permitido(90) == 12  # travado pelo maximo absoluto


def test_seccao_previsao_diz_quando_corta_o_horizonte():
    _, texto = gr.secao_previsao(serie_reta(12), 99)
    assert "máximo defensável" in texto
    assert "ficção" in texto


def test_seccao_previsao_explica_a_recusa_por_ajuste():
    serie = pd.Series([10, 90, 20, 85, 15, 95, 25, 80, 12, 88],
                      index=[f"P{i}" for i in range(10)])
    _, texto = gr.secao_previsao(serie, 3)
    assert "Não calculada" in texto
    assert "ajuste" in texto


def test_seccao_previsao_explica_a_recusa_por_serie_curta():
    _, texto = gr.secao_previsao(serie_reta(5), 2)
    assert "pelo menos 8" in texto


def test_previsao_avisa_quando_desce_abaixo_de_zero():
    serie = serie_reta(12, declive=-2.0) + 22  # comeca em 22 e desce ate zero
    _, texto = gr.secao_previsao(serie, 4)
    assert "abaixo de zero" in texto


def test_eixo_categorico_recusa_previsao():
    serie = pd.Series([10, 20, 30], index=["Meta", "Google", "TikTok"])
    blocos = gr.montar_analise(serie, grafico(eixo_x="Canal", previsao=3), 3, False, None)
    assert "período seguinte" in texto_de(blocos, "Previsão")


def test_sem_campo_previsao_nao_ha_seccao_previsao():
    serie = serie_reta(12)
    blocos = gr.montar_analise(serie, grafico(), 12, True, None)
    assert "Previsão" not in etiquetas(blocos)


def test_previsao_e_recusada_se_nao_for_inteiro_positivo():
    with pytest.raises(gr.ErroDados, match="previsao"):
        gr.validar_grafico(grafico(previsao=0), 1)
    with pytest.raises(gr.ErroDados, match="previsao"):
        gr.validar_grafico(grafico(previsao="seis"), 1)


# ------------------------------------------------------- rotulos futuros


def test_rotulos_futuros_viram_o_ano():
    rotulos = gr.proximos_rotulos(pd.Index(["2025-10", "2025-11", "2025-12"]), 3)
    assert rotulos == ["2026-01", "2026-02", "2026-03"]


def test_rotulos_futuros_continuam_os_meses():
    rotulos = gr.proximos_rotulos(pd.Index(["Outubro", "Novembro", "Dezembro"]), 2)
    assert rotulos == ["Janeiro", "Fevereiro"]


def test_rotulos_futuros_continuam_os_anos():
    assert gr.proximos_rotulos(pd.Index([2023, 2024, 2025]), 2) == ["2026", "2027"]


def test_rotulos_futuros_desistem_em_vez_de_inventar():
    rotulos = gr.proximos_rotulos(pd.Index(["Fase A", "Fase B"]), 2)
    assert rotulos == ["período +1", "período +2"]


# ------------------------------------------------------------ eixo temporal


def test_meses_em_portugues_sao_linha_do_tempo():
    assert gr.parece_temporal(pd.Index(MESES)) is True


def test_anos_sao_linha_do_tempo():
    assert gr.parece_temporal(pd.Index([2023, 2024, 2025])) is True


def test_canais_nao_sao_linha_do_tempo():
    """Uma regressao sobre «Canal» mudaria de declive so por trocar colunas."""
    assert gr.parece_temporal(pd.Index(["Meta", "Google Ads", "TikTok"])) is False


def test_extrair_ano_de_varios_formatos():
    assert gr.extrair_ano("2023-01") == 2023
    assert gr.extrair_ano(2024) == 2024
    assert gr.extrair_ano(pd.Timestamp("2025-06-30")) == 2025
    assert gr.extrair_ano("Janeiro") is None


def test_indice_sazonal_precisa_de_dois_ciclos():
    um_ano = pd.Series(range(1, 13), index=[f"2025-{m:02d}" for m in range(1, 13)])
    assert gr.indice_sazonal(um_ano) is None


def test_indice_sazonal_encontra_o_mes_forte():
    rotulos, valores = [], []
    for ano in (2024, 2025):
        for mes in range(1, 13):
            rotulos.append(f"{ano}-{mes:02d}")
            valores.append(300 if mes == 12 else 100)
    resultado = gr.indice_sazonal(pd.Series(valores, index=rotulos))

    assert resultado["ciclos"] == 2
    assert max(resultado["indices"], key=resultado["indices"].get) == 12


# -------------------------------------------------------------- bloco analise


def etiquetas(blocos):
    return [etiqueta for etiqueta, _ in blocos]


def texto_de(blocos, etiqueta):
    return next(t for e, t in blocos if e == etiqueta)


def test_eixo_categorico_nao_leva_tendencia():
    serie = pd.Series([10, 20, 30], index=["Meta", "Google", "TikTok"])
    blocos = gr.montar_analise(serie, grafico(eixo_x="Canal"), 3, False, None)

    assert "Tendência" not in etiquetas(blocos)
    assert "categórico" in texto_de(blocos, "Evolução")


def test_serie_curta_diz_que_a_regressao_nao_correu():
    serie = pd.Series([10, 20, 30], index=["Janeiro", "Fevereiro", "Marco"])
    blocos = gr.montar_analise(serie, grafico(), 3, True, None)

    assert "não calculada" in texto_de(blocos, "Tendência")
    assert "pelo menos 4" in texto_de(blocos, "Tendência")


def test_analise_ano_a_ano_com_dois_anos():
    anual = pd.Series([100.0, 150.0], index=[2024, 2025])
    serie = pd.Series([10, 20, 30, 40], index=MESES[:4])
    blocos = gr.montar_analise(serie, grafico(), 4, True, anual)

    comparacao = texto_de(blocos, "Comparação entre anos")
    assert "+50" in comparacao
    assert "+50%" in comparacao
    assert "2025" in texto_de(blocos, "Ano a ano")


def test_sem_juizos_de_valor_no_texto():
    """A linha que nao se atravessa: nada de «bom», «mau» ou «fraco desempenho»."""
    serie = pd.Series([10, 20, 15, 40, 35, 60], index=[f"2025-{m:02d}" for m in range(1, 7)])
    blocos = gr.montar_analise(serie, grafico(), 6, True, None)
    tudo = " ".join(t for _, t in blocos).lower()

    for proibida in ("bom ", "mau ", "excelente", "preocupante", "fraco desempenho"):
        assert proibida not in tudo, f"apareceu um juízo de valor: {proibida}"


# --------------------------------------------------------- bloqueio por avisos


def correr_cli(*argumentos: str) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, str(SCRIPT), *argumentos],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )


def test_avisos_bloqueiam_a_geracao_e_nao_criam_docx(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "sujo.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"], "Valor": [10, None, 30]},
    )
    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    saida = tmp_path / "relatorio.docx"

    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--saida", str(saida))

    assert resultado.returncode == 2
    assert not saida.exists(), "gerou o .docx apesar dos avisos"
    assert "--avisos-aceites" in resultado.stdout


def test_avisos_aceites_desbloqueiam_a_geracao(tmp_path):
    ficheiro = escrever_excel(
        tmp_path / "sujo.xlsx",
        {"Mes": ["Janeiro", "Fevereiro", "Marco"], "Valor": [10, None, 30]},
    )
    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    saida = tmp_path / "relatorio.docx"

    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--saida", str(saida), "--avisos-aceites")

    assert resultado.returncode == 0
    assert saida.exists()


def test_verificar_nunca_gera_nada(tmp_path):
    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES[:3], "Valor": [1, 2, 3]})
    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    saida = tmp_path / "relatorio.docx"

    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano),
                           "--saida", str(saida), "--verificar")

    assert resultado.returncode == 0
    assert not saida.exists()


def test_erro_de_dados_nao_mostra_traceback(tmp_path):
    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES[:3], "Valor": [1, 2, 3]})
    plano = escrever_plano(tmp_path / "plano.json", [grafico(folha="Inexistente")])

    resultado = correr_cli("--dados", str(ficheiro), "--plano", str(plano), "--verificar")

    assert resultado.returncode == 1
    assert "Traceback" not in resultado.stderr
    assert "Erro:" in resultado.stderr


# ------------------------------------------------------- limpeza e integridade


def test_o_excel_de_origem_nao_e_alterado(tmp_path):
    ficheiro = escrever_excel(tmp_path / "origem.xlsx",
                              {"Mes": MESES[:3], "Valor": [1, 2, 3]})
    antes = ficheiro.read_bytes()

    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    correr_cli("--dados", str(ficheiro), "--plano", str(plano),
               "--saida", str(tmp_path / "r.docx"))

    assert ficheiro.read_bytes() == antes


def test_nao_ficam_pastas_temporarias_para_tras(tmp_path):
    import tempfile

    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES[:3], "Valor": [1, 2, 3]})
    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    base = Path(tempfile.gettempdir())
    antes = set(base.glob("excel-para-word-*"))

    correr_cli("--dados", str(ficheiro), "--plano", str(plano),
               "--saida", str(tmp_path / "r.docx"))

    assert set(base.glob("excel-para-word-*")) == antes


def test_analise_curta_nao_traz_o_bloco_de_analise(tmp_path):
    from docx import Document

    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES, "Valor": list(range(1, 13))})
    plano = tmp_path / "plano.json"
    plano.write_text(json.dumps(
        {"titulo_relatorio": "T", "analise": "curta", "graficos": [grafico()]},
        ensure_ascii=False), encoding="utf-8")
    saida = tmp_path / "r.docx"
    correr_cli("--dados", str(ficheiro), "--plano", str(plano), "--saida", str(saida))

    texto = "\n".join(p.text for p in Document(str(saida)).paragraphs)
    assert "Análise" not in texto
    assert "O total é" in texto


def test_analise_completa_e_o_comportamento_por_omissao(tmp_path):
    from docx import Document

    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES, "Valor": list(range(1, 13))})
    plano = escrever_plano(tmp_path / "plano.json", [grafico()])
    saida = tmp_path / "r.docx"
    correr_cli("--dados", str(ficheiro), "--plano", str(plano), "--saida", str(saida))

    texto = "\n".join(p.text for p in Document(str(saida)).paragraphs)
    assert "Análise" in texto
    assert "Dispersão" in texto


def test_campo_desconhecido_no_plano_e_recusado(tmp_path):
    plano = tmp_path / "plano.json"
    plano.write_text(json.dumps(
        {"titulo_relatorio": "T", "graficos": [grafico()], "autor": "eu"},
        ensure_ascii=False), encoding="utf-8")
    with pytest.raises(gr.ErroDados, match="autor"):
        gr.ler_plano(plano)


def test_analise_com_valor_invalido_e_recusada(tmp_path):
    plano = tmp_path / "plano.json"
    plano.write_text(json.dumps(
        {"titulo_relatorio": "T", "analise": "profunda", "graficos": [grafico()]},
        ensure_ascii=False), encoding="utf-8")
    with pytest.raises(gr.ErroDados, match="profunda"):
        gr.ler_plano(plano)


def test_documento_gerado_tem_o_conteudo_esperado(tmp_path):
    from docx import Document

    ficheiro = escrever_excel(tmp_path / "limpo.xlsx",
                              {"Mes": MESES[:3], "Valor": [10, 20, 30]})
    plano = escrever_plano(
        tmp_path / "plano.json",
        [grafico(titulo="Vendas por mês", tabela_dados=True)],
        titulo="Relatório anual",
    )
    saida = tmp_path / "r.docx"
    correr_cli("--dados", str(ficheiro), "--plano", str(plano), "--saida", str(saida))

    documento = Document(str(saida))
    texto = "\n".join(p.text for p in documento.paragraphs)

    assert "Relatório anual" in texto
    assert "Vendas por mês" in texto
    assert "Notas sobre os dados" in texto
    assert "O total é 60." in texto
    assert len(documento.inline_shapes) == 1, "faltou a imagem do gráfico"
    assert len(documento.tables) == 1
    assert documento.tables[0].rows[1].cells[0].text == "Janeiro"
